"""Generate Word (.docx) reports from classified articles.

Uses python-docx to create professional reports with sentiment-based sections.
Falls back to a simple text-based approach if python-docx is not installed.
"""
from __future__ import annotations

import logging
from datetime import datetime, timezone
from pathlib import Path

from news_monitor.models import CycleReport, NewsArticle

logger = logging.getLogger("news_monitor.output.docx_reporter")

# Sentiment categories for event-based monitoring
_SENTIMENT_CATEGORIES = [
    ("positive", "正面/建设性报道"),
    ("neutral",  "中性/客观报道"),
    ("negative", "负面/批评报道"),
    ("analysis", "深度分析"),
]

# Colors for sentiment labels (RGB tuples)
_SENTIMENT_COLORS = {
    "positive": (34, 139, 34),    # forest green
    "neutral":  (70, 130, 180),   # steel blue
    "negative": (205, 92, 92),    # indian red
    "analysis": (148, 103, 189),  # purple
}


def _trunc(text: str, n: int) -> str:
    return text[:n] + "..." if len(text) > n else text


def _group_by_sentiment(
    articles: list[NewsArticle],
) -> dict[str, list[NewsArticle]]:
    """Group articles by sentiment category."""
    groups: dict[str, list[NewsArticle]] = {cat: [] for cat, _ in _SENTIMENT_CATEGORIES}
    for art in articles:
        cat = art.category if art.category in groups else "neutral"
        groups[cat].append(art)
    return groups


def generate_docx_report(
    articles: list[NewsArticle],
    report: CycleReport,
    output_dir: str,
    event_name: str,
    sentiment_log: list[dict] | None = None,
    ts: str | None = None,
    executive_summary: str = "",
) -> str:
    """Generate a Word document report. Returns the file path.

    Requires python-docx. If not available, raises ImportError.
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if ts is None:
        ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    docx_path = str(Path(output_dir) / f"{ts}_boao_forum.docx")

    doc = Document()

    # -- Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # -- Title
    title_para = doc.add_heading(f"{event_name} — 海外媒体舆情报告", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -- Generation info
    gen_time = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(f"生成时间：{gen_time}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # -- Executive summary
    if executive_summary:
        doc.add_heading("综合舆情摘要", level=1)
        sum_para = doc.add_paragraph(executive_summary)
        sum_para.paragraph_format.left_indent = Cm(0.5)
        doc.add_paragraph()

    # -- Summary table
    doc.add_heading("数据概览", level=1)
    groups = _group_by_sentiment(articles)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "指标"
    hdr[1].text = "数量"

    stats = [
        ("抓取文章数", str(report.total_fetched)),
        ("去重后数量", str(report.total_after_dedup)),
        ("已翻译数量", str(report.total_translated)),
        ("最终收录", str(len(articles))),
    ]
    for cat_key, cat_label in _SENTIMENT_CATEGORIES:
        n = len(groups[cat_key])
        if n:
            stats.append((f"  └ {cat_label}", str(n)))

    for label, value in stats:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    doc.add_paragraph()

    # -- Sentiment distribution overview
    doc.add_heading("舆情态度分布", level=1)
    dist_text = []
    for cat_key, cat_label in _SENTIMENT_CATEGORIES:
        n = len(groups[cat_key])
        if n:
            pct = n / max(len(articles), 1) * 100
            dist_text.append(f"{cat_label}：{n} 篇（{pct:.0f}%）")
    doc.add_paragraph("  |  ".join(dist_text))
    doc.add_paragraph()

    # -- Importance highlights (top articles)
    if sentiment_log:
        high_importance = [e for e in sentiment_log if e.get("importance", 0) >= 4]
        if high_importance:
            high_importance.sort(key=lambda e: e.get("importance", 0), reverse=True)
            doc.add_heading("重要舆情摘要（重要度 ≥ 4）", level=1)
            for entry in high_importance[:15]:
                title = entry.get("title", "")
                attitude = entry.get("key_attitude", "")
                sentiment = entry.get("sentiment", "neutral")
                importance = entry.get("importance", 3)
                color = _SENTIMENT_COLORS.get(sentiment, (0, 0, 0))

                p = doc.add_paragraph()
                # Importance badge
                badge_run = p.add_run(f"[★{importance}] ")
                badge_run.font.bold = True
                badge_run.font.size = Pt(10)
                # Sentiment tag
                tag_run = p.add_run(f"[{sentiment}] ")
                tag_run.font.color.rgb = RGBColor(*color)
                tag_run.font.size = Pt(9)
                # Title
                title_run = p.add_run(_trunc(title, 80))
                title_run.font.bold = True
                title_run.font.size = Pt(10)
                # Attitude on next line
                if attitude:
                    att_p = doc.add_paragraph()
                    att_p.paragraph_format.left_indent = Cm(1)
                    att_run = att_p.add_run(f"→ {attitude}")
                    att_run.font.size = Pt(9)
                    att_run.font.color.rgb = RGBColor(80, 80, 80)

            doc.add_paragraph()

    # -- Sections by sentiment
    for cat_key, cat_label in _SENTIMENT_CATEGORIES:
        arts = groups[cat_key]
        if not arts:
            continue

        color = _SENTIMENT_COLORS.get(cat_key, (0, 0, 0))

        heading = doc.add_heading(f"■ {cat_label}（{len(arts)} 篇）", level=1)
        # Color the heading
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)

        # Build importance lookup from log
        importance_map: dict[str, int] = {}
        attitude_map: dict[str, str] = {}
        if sentiment_log:
            for entry in sentiment_log:
                t = entry.get("title", "")
                importance_map[t] = entry.get("importance", 3)
                attitude_map[t] = entry.get("key_attitude", "")

        # Sort by importance descending
        arts.sort(key=lambda a: importance_map.get(a.title_zh or a.title, 3), reverse=True)

        for idx, art in enumerate(arts, 1):
            title = art.title_zh or art.title
            summary = art.summary_zh or art.description_zh or art.description or ""
            importance = importance_map.get(title, 3)

            # Article title
            p = doc.add_heading(f"{idx}. {_trunc(title, 80)}", level=2)

            # Key attitude
            attitude = attitude_map.get(title, "")
            if attitude:
                att_p = doc.add_paragraph()
                att_run = att_p.add_run(f"核心态度：{attitude}")
                att_run.font.bold = True
                att_run.font.size = Pt(10)
                att_run.font.color.rgb = RGBColor(*color)

            # Summary
            if summary:
                sum_p = doc.add_paragraph(_trunc(summary, 400))
                sum_p.paragraph_format.left_indent = Cm(0.5)

            # Metadata
            meta_parts = []
            if art.source_name:
                meta_parts.append(f"来源：{art.source_name}")
            if art.found_by:
                meta_parts.append(f"搜索引擎：{'+'.join(art.found_by)}")
            if art.published_at:
                meta_parts.append(f"发布：{art.published_at.strftime('%Y-%m-%d %H:%M')}")
            meta_parts.append(f"重要度：{'★' * importance}")

            meta_p = doc.add_paragraph()
            meta_run = meta_p.add_run("  |  ".join(meta_parts))
            meta_run.font.size = Pt(8)
            meta_run.font.color.rgb = RGBColor(128, 128, 128)

            if art.source_url:
                link_p = doc.add_paragraph()
                link_run = link_p.add_run(art.source_url)
                link_run.font.size = Pt(8)
                link_run.font.color.rgb = RGBColor(30, 100, 200)

    # Save
    doc.save(docx_path)
    logger.info("DOCX written: %s  (%d articles)", docx_path, len(articles))
    return docx_path
