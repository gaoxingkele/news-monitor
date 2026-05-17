"""Convert v2.0 markdown to DOCX with proper Chinese formatting."""
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

MD_FILE = r'D:\aicoding\news-monitor\topics\0513\台湾同温层研究报告_v2.1.md'
DOCX_FILE = r'D:\aicoding\news-monitor\topics\0513\台湾同温层研究报告_v2.1.docx'

CN_FONT = '宋体'
HEADING_FONT = '黑体'

def set_cn_font(run, font=CN_FONT, size_pt=12, bold=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    run.font.size = Pt(size_pt)
    run.bold = bold

def add_heading_para(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    sizes = {1: 20, 2: 16, 3: 14, 4: 13, 5: 12, 6: 12}
    r = p.add_run(text)
    set_cn_font(r, font=HEADING_FONT, size_pt=sizes.get(level, 12), bold=True)
    return p

def process_inline_bold(p, text, font=CN_FONT, size_pt=11):
    """Split text by **...** and produce alternating bold/normal runs."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            set_cn_font(r, font=font, size_pt=size_pt, bold=True)
        else:
            r = p.add_run(part)
            set_cn_font(r, font=font, size_pt=size_pt, bold=False)


def add_table_from_md(doc, table_lines):
    """Convert pipe-delimited markdown lines to a Word table."""
    rows = []
    for line in table_lines:
        if not line.strip():
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        # Skip the alignment separator row
        if all(re.match(r'^[-:\s]+$', c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < n_cols:
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                # process bold
                process_inline_bold(p, cell_text, size_pt=10)
                if i == 0:
                    for run in p.runs:
                        run.bold = True


def md_to_docx():
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Set normal style font
    style = doc.styles['Normal']
    style.font.name = CN_FONT
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), CN_FONT)

    i = 0
    while i < len(lines):
        line = lines[i]
        # Tables: detect with `|`
        if line.startswith('|'):
            tbl = []
            while i < len(lines) and lines[i].startswith('|'):
                tbl.append(lines[i])
                i += 1
            add_table_from_md(doc, tbl)
            continue
        # Headings
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            add_heading_para(doc, m.group(2).strip(), level)
            i += 1
            continue
        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            r = p.add_run('—' * 30)
            set_cn_font(r, size_pt=10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        # Blockquote
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            text = line[2:]
            process_inline_bold(p, text)
            i += 1
            continue
        # Bullet
        m = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if m:
            indent = len(m.group(1))
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.75 + indent * 0.4)
            process_inline_bold(p, m.group(2))
            i += 1
            continue
        # Numbered list
        m = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
        if m:
            indent = len(m.group(1))
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75 + indent * 0.4)
            process_inline_bold(p, f'{m.group(2)}. {m.group(3)}')
            i += 1
            continue
        # Empty line
        if not line.strip():
            i += 1
            continue
        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.75)
        p.paragraph_format.line_spacing = 1.5
        process_inline_bold(p, line)
        i += 1

    doc.save(DOCX_FILE)
    print(f'Saved: {DOCX_FILE}')
    import os
    print(f'Size: {os.path.getsize(DOCX_FILE)} bytes')


if __name__ == '__main__':
    md_to_docx()
