"""米切尔研究所 v2 报告 markdown → docx 转换。

中文期刊论文格式：仿宋四号 / 1.5行距 / 首行缩进 / 黑体标题。
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(r"D:\aicoding\news-monitor\topics\0430\米切尔研究所_深度专家分析_v2.md")
OUT = SRC.with_suffix(".docx")

BODY_FONT = "仿宋"
HEAD_FONT = "黑体"
MONO_FONT = "Consolas"

BODY_SIZE = Pt(12)
HEAD1_SIZE = Pt(20)
HEAD2_SIZE = Pt(16)
HEAD3_SIZE = Pt(14)
HEAD4_SIZE = Pt(13)
LINE_SPACING = 1.5
INDENT_EM = Cm(0.74)


def set_run_font(run, face, size, bold=False, italic=False, color=None):
    run.bold = bold
    run.italic = italic
    run.font.name = face
    run.font.size = size
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), face)
    rPr.insert(0, rFonts)


def add_para(doc, text, face=BODY_FONT, size=BODY_SIZE, bold=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=Pt(0),
             space_after=Pt(4), indent=None, line_spacing=LINE_SPACING,
             color=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if indent is not None:
        pf.first_line_indent = indent
    add_runs_with_inline(p, text, face, size, bold, color)
    return p


# ── inline parser: **bold**, `code`, [text](url) ────────────────────────────
INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)


def add_runs_with_inline(p, text, face, size, bold_default=False, color=None):
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            set_run_font(r, face, size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            set_run_font(r, MONO_FONT, size, bold=bold_default, color=(80, 80, 80))
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                txt, url = m.group(1), m.group(2)
                r = p.add_run(f"{txt} ({url})")
                set_run_font(r, face, size, bold=bold_default, color=(0, 0, 180))
            else:
                r = p.add_run(part)
                set_run_font(r, face, size, bold=bold_default, color=color)
        else:
            r = p.add_run(part)
            set_run_font(r, face, size, bold=bold_default, color=color)


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    # 默认正文字体
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)

    in_code = False
    code_buf = []
    in_blockquote = False

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # code fence
        if line.startswith("```"):
            in_code = not in_code
            if not in_code and code_buf:
                add_para(doc, "\n".join(code_buf), face=MONO_FONT,
                         size=Pt(10), color=(60, 60, 60))
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # h1
        if line.startswith("# "):
            add_para(doc, line[2:].strip(), face=HEAD_FONT, size=HEAD1_SIZE,
                     bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=Pt(12), space_after=Pt(14))
        elif line.startswith("## "):
            add_para(doc, line[3:].strip(), face=HEAD_FONT, size=HEAD2_SIZE,
                     bold=True, space_before=Pt(16), space_after=Pt(8))
        elif line.startswith("### "):
            add_para(doc, line[4:].strip(), face=HEAD_FONT, size=HEAD3_SIZE,
                     bold=True, space_before=Pt(10), space_after=Pt(6))
        elif line.startswith("#### "):
            add_para(doc, line[5:].strip(), face=HEAD_FONT, size=HEAD4_SIZE,
                     bold=True, space_before=Pt(8), space_after=Pt(4))
        elif line.startswith("> "):
            add_para(doc, line[2:].strip(), face=BODY_FONT, size=Pt(11),
                     color=(90, 90, 90), indent=Cm(0.5))
        elif re.match(r"^[-\*]\s", line):
            txt = re.sub(r"^[-\*]\s", "", line).strip()
            add_para(doc, "• " + txt, indent=Cm(0.5), space_after=Pt(2))
        elif re.match(r"^\d+\.\s", line):
            add_para(doc, line.strip(), indent=Cm(0.5), space_after=Pt(2))
        elif re.match(r"^---+\s*$", line):
            add_para(doc, "─" * 30, align=WD_ALIGN_PARAGRAPH.CENTER,
                     color=(150, 150, 150), space_before=Pt(6),
                     space_after=Pt(6))
        elif line.strip() == "":
            # 空行：跳过
            pass
        else:
            add_para(doc, line, indent=INDENT_EM)
        i += 1

    doc.save(OUT)
    sz = OUT.stat().st_size
    print(f"[OK] 写入 {OUT} ({sz:,} bytes)")


if __name__ == "__main__":
    main()
