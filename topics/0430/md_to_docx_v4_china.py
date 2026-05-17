"""V4 中国智库风格 markdown → docx 转换器。

排版规范（参考中国政府公文 GB/T 9704-2012 + 主流智库报告习惯）：
- 标题：方正小标宋 / 黑体（一级居中加粗）
- 正文：仿宋_GB2312 小四号（12pt）
- 行距：1.5 倍
- 段落首行缩进 2 字符
- 引用编号 [N] → 上标
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(r"D:\aicoding\news-monitor\topics\0430\米切尔研究所_战略研判_v4.md")
OUT = SRC.with_suffix(".docx")

BODY_FONT_EN = "Times New Roman"
BODY_FONT_CN = "仿宋"
HEAD_FONT_CN = "黑体"
TITLE_FONT_CN = "方正小标宋简体"  # 政府公文标题字体；缺失时回退到黑体
MONO_FONT = "Consolas"

BODY_SIZE = Pt(12)         # 小四号
TITLE_SIZE = Pt(22)        # 一号字（标题）
H1_SIZE = Pt(18)           # 小二（一级）
H2_SIZE = Pt(15)           # 三号（二级）
H3_SIZE = Pt(13)           # 小三（三级）
LINE_SPACING = 1.5
INDENT_2CHAR = Cm(0.74)    # 首行缩进 2 字符（小四 ≈ 0.74cm）


def set_font(run, en_face=BODY_FONT_EN, cn_face=BODY_FONT_CN, size=BODY_SIZE,
             bold=False, italic=False, color=None, superscript=False):
    run.bold = bold
    run.italic = italic
    if superscript:
        run.font.superscript = True
    run.font.name = en_face
    run.font.size = size
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), cn_face)
    rPr.insert(0, rFonts)


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|\[\^?\d{1,3}\])")


def add_inline(p, text, en_face=BODY_FONT_EN, cn_face=BODY_FONT_CN,
               size=BODY_SIZE, bold_default=False, color=None):
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            set_font(r, en_face, HEAD_FONT_CN, size, bold=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            set_font(r, MONO_FONT, MONO_FONT, size, color=(80, 80, 80))
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                txt, url = m.group(1), m.group(2)
                r = p.add_run(txt)
                set_font(r, en_face, cn_face, size, color=(0, 0, 180))
                r2 = p.add_run(f" ({url})")
                set_font(r2, en_face, cn_face, Pt(9), color=(100, 100, 180))
        elif re.match(r"^\[\^?\d{1,3}\]$", part):
            r = p.add_run(part)
            set_font(r, en_face, cn_face, Pt(8), color=(50, 50, 150),
                     superscript=True)
        else:
            r = p.add_run(part)
            set_font(r, en_face, cn_face, size, bold=bold_default, color=color)


def add_para(doc, text, en_face=BODY_FONT_EN, cn_face=BODY_FONT_CN,
             size=BODY_SIZE, bold=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=Pt(0),
             space_after=Pt(4), indent=None, line_spacing=LINE_SPACING,
             color=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if indent is not None:
        pf.first_line_indent = indent
    add_inline(p, text, en_face, cn_face, size, bold, color)
    return p


def add_table_from_md(doc, rows: list[list[str]]):
    if not rows or len(rows) < 2:
        return
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Light Grid"
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            c = tbl.cell(ri, ci)
            c.text = ""
            p = c.paragraphs[0]
            add_inline(p, cell.strip(), size=Pt(10),
                       cn_face=HEAD_FONT_CN if ri == 0 else BODY_FONT_CN,
                       bold_default=(ri == 0))


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.8)

    # 默认正文
    style = doc.styles['Normal']
    style.font.name = BODY_FONT_EN
    style.font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT_CN)

    in_code = False
    code_buf: list[str] = []
    table_buf: list[list[str]] = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_table_from_md(doc, table_buf)
            doc.add_paragraph()
            table_buf = []

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 代码块
        if line.startswith("```"):
            in_code = not in_code
            if not in_code and code_buf:
                add_para(doc, "\n".join(code_buf), en_face=MONO_FONT,
                         cn_face=MONO_FONT, size=Pt(9), color=(60, 60, 60))
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 表格
        if "|" in line and line.count("|") >= 2 and not line.strip().startswith("|---"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(re.match(r"^[-:\s]*$", c) for c in cells):
                i += 1
                continue
            table_buf.append(cells)
            i += 1
            continue
        elif "|---" in line or re.match(r"^\|[-:\s|]+\|$", line.strip()):
            i += 1
            continue
        else:
            flush_table()

        # 标题
        if line.startswith("# "):
            add_para(doc, line[2:].strip(), en_face=HEAD_FONT_CN,
                     cn_face=TITLE_FONT_CN, size=TITLE_SIZE, bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=Pt(20), space_after=Pt(16),
                     color=(20, 30, 80))
        elif line.startswith("## "):
            add_para(doc, line[3:].strip(), en_face=HEAD_FONT_CN,
                     cn_face=HEAD_FONT_CN, size=H1_SIZE, bold=True,
                     space_before=Pt(20), space_after=Pt(10),
                     color=(20, 30, 80))
        elif line.startswith("### "):
            add_para(doc, line[4:].strip(), en_face=HEAD_FONT_CN,
                     cn_face=HEAD_FONT_CN, size=H2_SIZE, bold=True,
                     space_before=Pt(14), space_after=Pt(6),
                     color=(40, 50, 110))
        elif line.startswith("#### "):
            add_para(doc, line[5:].strip(), en_face=HEAD_FONT_CN,
                     cn_face=HEAD_FONT_CN, size=H3_SIZE, bold=True,
                     space_before=Pt(10), space_after=Pt(4),
                     color=(60, 70, 130))
        elif re.match(r"^[-*]{3,}\s*$", line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("─" * 25)
            set_font(r, color=(160, 160, 160))
        elif re.match(r"^>\s", line):
            # 引用块 — 用斜体 + 缩进
            content = line.lstrip("> ")
            p = add_para(doc, content, size=Pt(11), color=(80, 80, 80),
                         indent=Cm(0.74), space_after=Pt(2))
        elif re.match(r"^\s*[-\*]\s", line):
            txt = re.sub(r"^\s*[-\*]\s", "", line)
            add_para(doc, "•　" + txt, indent=Cm(0.5),
                     space_after=Pt(2), size=Pt(11.5))
        elif re.match(r"^\s*\d+\.\s", line) or re.match(r"^\s*\(\s*[一二三四五六七八九十]+\s*\)", line):
            add_para(doc, line.strip(), indent=Cm(0.5),
                     space_after=Pt(2), size=Pt(11.5))
        elif line.strip() == "":
            pass
        else:
            # 正文：首行缩进 2 字符
            add_para(doc, line, indent=INDENT_2CHAR, space_after=Pt(4))
        i += 1

    flush_table()

    # 页脚
    sec = doc.sections[0]
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("MISJ-2026-04 · 战略研判 · 内部参考")
    set_font(r, size=Pt(8), color=(120, 120, 120))

    doc.save(OUT)
    sz = OUT.stat().st_size
    print(f"[OK] 写入 {OUT} ({sz:,} bytes)")


if __name__ == "__main__":
    main()
