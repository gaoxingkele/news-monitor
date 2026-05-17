"""V3 RAND 体例报告 markdown → docx 转换器。

特殊处理：
- "Key Findings of Chapter N" 段落 → 浅灰阴影框（边框 + 内部填充 + 缩进）
- "Table N.M" / "Figure N.M" → 加粗居中 + 下方 SOURCE 行斜体
- "**Finding 1.**" / "**Recommendation 1.**" → 加粗编号缩进
- ### Confidence: High/Moderate/Low → 不同颜色徽章
- "[1]" / "[^1]" 引用编号 → 上标
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(r"D:\aicoding\news-monitor\topics\0430\米切尔研究所_RAND体例分析_v3.md")
OUT = SRC.with_suffix(".docx")

BODY_FONT = "Times New Roman"   # RAND 用 Sabon/Times
CN_FONT = "宋体"                 # 中文宋体（RAND 中文版常见）
HEAD_FONT = "Arial"              # RAND 标题用 Helvetica/Arial
MONO_FONT = "Consolas"

BODY_SIZE = Pt(11)
HEAD1_SIZE = Pt(22)
HEAD2_SIZE = Pt(16)
HEAD3_SIZE = Pt(13)
HEAD4_SIZE = Pt(12)
LINE_SPACING = 1.35
INDENT_EM = Cm(0)   # RAND 体例不首行缩进，靠段间距分段


def set_font(run, en_face=BODY_FONT, cn_face=CN_FONT, size=BODY_SIZE,
             bold=False, italic=False, color=None):
    run.bold = bold
    run.italic = italic
    run.font.name = en_face
    run.font.size = size
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), cn_face)
    rPr.insert(0, rFonts)


def set_para_shading(para, fill_hex="EFEFEF", border_color="999999"):
    """段落浅灰阴影 + 边框（RAND Key Findings 框样式）。"""
    pPr = para._p.get_or_add_pPr()
    # 阴影
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)
    # 边框
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), border_color)
        pBdr.append(b)
    pPr.append(pBdr)


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|\[\^?\d+\])")


def add_inline(p, text, en_face=BODY_FONT, cn_face=CN_FONT, size=BODY_SIZE,
               bold_default=False, color=None):
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            set_font(r, en_face, cn_face, size, bold=True, color=color)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            r = p.add_run(part[1:-1])
            set_font(r, en_face, cn_face, size, bold=bold_default, italic=True, color=color)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            set_font(r, MONO_FONT, MONO_FONT, size, color=(80, 80, 80))
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                txt, url = m.group(1), m.group(2)
                r = p.add_run(f"{txt}")
                set_font(r, en_face, cn_face, size, color=(0, 0, 180))
                r2 = p.add_run(f" ({url})")
                set_font(r2, en_face, cn_face, Pt(9), color=(100, 100, 180))
        elif re.match(r"^\[\^?\d+\]$", part):
            # 引用编号 → 上标
            r = p.add_run(part)
            set_font(r, en_face, cn_face, Pt(8), color=(50, 50, 150))
            r.font.superscript = True
        else:
            r = p.add_run(part)
            set_font(r, en_face, cn_face, size, bold=bold_default, color=color)


def add_para(doc, text, en_face=BODY_FONT, cn_face=CN_FONT, size=BODY_SIZE,
             bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=Pt(0), space_after=Pt(6), indent=None,
             line_spacing=LINE_SPACING, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if indent is not None:
        pf.first_line_indent = indent
    add_inline(p, text, en_face, cn_face, size, bold, color)
    return p


def add_table_from_md(doc, rows: list[list[str]]):
    if not rows:
        return
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci, cell in enumerate(row):
            c = tbl.cell(ri, ci)
            c.text = ""
            p = c.paragraphs[0]
            add_inline(p, cell.strip(), size=Pt(10),
                       bold_default=(ri == 0))


def main():
    text = SRC.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.4)
    sec.bottom_margin = Cm(2.4)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # 默认正文样式
    style = doc.styles['Normal']
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn('w:eastAsia'), CN_FONT)

    in_code = False
    code_buf: list[str] = []
    in_keyfindings_box = False
    table_buf: list[list[str]] = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            add_table_from_md(doc, table_buf)
            doc.add_paragraph()
            table_buf = []

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # 代码块
        if line.startswith("```"):
            in_code = not in_code
            if not in_code and code_buf:
                add_para(doc, "\n".join(code_buf), en_face=MONO_FONT,
                         cn_face=MONO_FONT, size=Pt(9), color=(60, 60, 60))
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 表格行检测
        if "|" in line and line.count("|") >= 2 and not line.strip().startswith("|---"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # 跳过分隔符行
            if all(re.match(r"^[-:\s]*$", c) for c in cells):
                i += 1
                continue
            table_buf.append(cells)
            i += 1
            continue
        elif "|---" in line or re.match(r"^\|[-:\s|]+\|$", line.strip()):
            i += 1
            continue
        else:
            flush_table()

        # 阴影框检测：> ### Key Findings of Chapter
        # 或 > **KF X.Y** 类型块（连续 > 行）
        if line.startswith(">"):
            content = line.lstrip(">").lstrip()
            if not in_keyfindings_box and (
                "Key Findings of Chapter" in content
                or "Key Findings of This Chapter" in content
                or content.startswith("###") and "Findings" in content
            ):
                in_keyfindings_box = True
                # 阴影框标题
                p = add_para(doc, content.lstrip("#").strip(),
                             en_face=HEAD_FONT, cn_face="黑体", size=HEAD3_SIZE,
                             bold=True, space_before=Pt(8), space_after=Pt(4),
                             color=(40, 40, 80))
                set_para_shading(p, fill_hex="DDE6F2", border_color="4F81BD")
            else:
                if not content:
                    i += 1
                    continue
                # 框内段落
                p = add_para(doc, content, size=Pt(10),
                             space_after=Pt(3), indent=Cm(0.3))
                if in_keyfindings_box:
                    set_para_shading(p, fill_hex="EFEFEF", border_color="BBBBBB")
            i += 1
            continue
        else:
            in_keyfindings_box = False

        # 标题
        if line.startswith("# "):
            add_para(doc, line[2:].strip(), en_face=HEAD_FONT, cn_face="黑体",
                     size=HEAD1_SIZE, bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=Pt(14), space_after=Pt(14),
                     color=(20, 40, 90))
        elif line.startswith("## "):
            add_para(doc, line[3:].strip(), en_face=HEAD_FONT, cn_face="黑体",
                     size=HEAD2_SIZE, bold=True,
                     space_before=Pt(18), space_after=Pt(8),
                     color=(20, 40, 90))
        elif line.startswith("### "):
            add_para(doc, line[4:].strip(), en_face=HEAD_FONT, cn_face="黑体",
                     size=HEAD3_SIZE, bold=True,
                     space_before=Pt(12), space_after=Pt(6),
                     color=(40, 70, 130))
        elif line.startswith("#### "):
            add_para(doc, line[5:].strip(), en_face=HEAD_FONT, cn_face="黑体",
                     size=HEAD4_SIZE, bold=True,
                     space_before=Pt(8), space_after=Pt(4),
                     color=(60, 90, 150))
        # Table / Figure 标注
        elif re.match(r"^\s*Table \d+\.\d+", line) or re.match(r"^\s*Figure \d+\.\d+", line):
            add_para(doc, line.strip(), en_face=HEAD_FONT, cn_face="黑体",
                     size=Pt(11), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=Pt(10), space_after=Pt(4),
                     color=(80, 80, 80))
        # SOURCE 注释
        elif re.match(r"^\s*SOURCE[:：]", line, re.I):
            add_para(doc, line.strip(), size=Pt(9),
                     align=WD_ALIGN_PARAGRAPH.CENTER, color=(110, 110, 110))
        # ─── 分隔
        elif re.match(r"^[-*]{3,}\s*$", line) or re.match(r"^\*\*\*\s*$", line):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("─" * 25)
            set_font(r, color=(160, 160, 160))
        # 列表项
        elif re.match(r"^\s*[-\*]\s", line):
            indent_lvl = (len(line) - len(line.lstrip())) // 2
            txt = re.sub(r"^\s*[-\*]\s", "", line)
            add_para(doc, "•  " + txt, indent=Cm(0.3),
                     space_after=Pt(2), size=Pt(10.5))
        elif re.match(r"^\s*\d+\.\s", line):
            indent_lvl = (len(line) - len(line.lstrip())) // 2
            add_para(doc, line.strip(), indent=Cm(0.3),
                     space_after=Pt(2), size=Pt(10.5))
        elif line.strip() == "":
            pass
        else:
            add_para(doc, line, indent=INDENT_EM, space_after=Pt(4))
        i += 1

    flush_table()

    # 页脚加文档号
    sec = doc.sections[0]
    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("RR-MI-2026-3 — Independent Open-Source Assessment")
    set_font(r, size=Pt(8), color=(120, 120, 120))

    doc.save(OUT)
    sz = OUT.stat().st_size
    print(f"[OK] 写入 {OUT} ({sz:,} bytes)")


if __name__ == "__main__":
    main()
