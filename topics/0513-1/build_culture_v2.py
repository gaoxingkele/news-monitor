# -*- coding: utf-8 -*-
"""Build Culture Volume v2.0 — Lai Ching-te era cognitive warfare deep analysis."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\aicoding\news-monitor\topics\0513-1\台湾认知作战与同温层固化深度分析报告_文化卷_v2.0.docx"

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for s_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    s = doc.styles[s_name]
    s.font.name = 'Microsoft YaHei'
    s.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    s.font.color.rgb = RGBColor(0x6B, 0x1F, 0x3A)

for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)


def set_cell_bg(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_p(text, bold=False, size=11, align=None, color=None, italic=False, space_after=4):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.35
    return p


def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        elif level == 3:
            run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(0x6B, 0x1F, 0x3A)
    return h


def add_table(headers, rows, header_color='6B1F3A', highlight_col=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 2'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        set_cell_bg(hdr[i], header_color)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Microsoft YaHei'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if highlight_col and ci == highlight_col:
                set_cell_bg(cells[ci], 'FCEEE6')
    return t


def add_blockquote(text, color=(0x6B, 0x1F, 0x3A)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), '6B1F3A')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(*color)
    run.italic = True
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return p


# ============ Cover ============
add_p('赖清德上台以来台湾岛内"同温层"固化与议题破圈规律深度研究',
      bold=True, size=20, align='center', color=(0x6B, 0x1F, 0x3A), space_after=4)
add_p('——基于认知作战事件谱、文化教育政策与舆论生态的多维分析（文化卷 v2.0）',
      size=12, align='center', color=(0x55, 0x55, 0x55), italic=True, space_after=14)

add_p('姊妹篇：与《台湾经济结构性脆弱深度分析报告 v2.0》（经济卷）配套，共同构成'
      '"赖清德执政期台湾结构性变化"双卷研究。',
      size=10, align='center', color=(0x66, 0x66, 0x66))
add_p('研究方法：通过 Brave / Tavily / Gemini / Grok 类公开检索，整合台湾文化部、教育部、'
      '陆委会公开资料；CNA中央社、自由时报、联合报、中时、Newtalk、RFA、Taipei Times、'
      'Wikipedia；政大选研中心、台湾民意基金会、TVBS民调中心；新华社、人民日报、'
      '国台办、华夏经纬、观察者网、福建日报等多源数据。',
      size=10, align='center', color=(0x66, 0x66, 0x66), space_after=10)
add_p('引用规范：APA 修订格式，正文以 [编号] 标注，文末附完整 References。',
      size=10, align='center', color=(0x66, 0x66, 0x66), space_after=18)

# ============ 引言：同温层概念 ============
add_h('引言：什么是"同温层"？——为非专业读者的简要说明', level=1)
add_p('"同温层"（filter bubble / echo chamber，又译"过滤气泡""回声室"）是政治传播学的核心概念之一。'
      '原本指大气层中温度相对均匀、空气流动受阻的层域；引申到社群媒体时代，'
      '指演算法（algorithm）依据使用者过去的点赞、停留、转发行为，反复推送其偏好的信息，'
      '使个人长期暴露于观点相近的内容，逐渐与异质观点隔离的现象 [1][2]。', align='justify')

add_p('同温层的三个学理层次：', bold=True)
add_p('① 过滤气泡（Pariser, 2011）：技术层——演算法依据用户画像过滤异己内容，造成信息环境同质化 [1]。',
      align='justify')
add_p('② 回声室（Sunstein, 2001）：心理层——用户主动选择志同道合的圈子，相同观点反复回响、加强 [2]。',
      align='justify')
add_p('③ 信息茧房（Sunstein, 2006）：行为层——长期处于过滤+回声环境后，认知边界自我封闭，'
      '不再接收外部异议，群体极化（group polarization）加剧 [2][3]。', align='justify')

add_blockquote('对本报告的意义：当一个社会的舆论场被多个相互不通气的"同温层"分割，'
               '主流议题在层与层之间难以"破圈"（即跨越同温层边界、形成共识）。'
               '执政者若善于利用并强化这种结构，就能在不必说服反对者的情况下，'
               '靠动员己方同温层来维持治理；而反对者的声音则被困在自己的同温层中，'
               '无法形成对执政者的有效制约。这就是赖清德执政以来台湾舆论场的核心特征。')

# ============ I 同温层固化 ============
add_h('一、赖清德执政期台湾"同温层"固化的新态势', level=1)

add_h('1.1 同温层固化的三大量化指标', level=2)
add_table(
    ['指标', '数据（2024–2026）', '同温层意义', '来源'],
    [
        ['"自认台湾人"比例', '61.7%（2024/2）→ 63.4%（2024/12）→ 略增（2025/7）[4][5]',
         '主流叙事（台湾主体性）已固化为多数共识', '政大选研'],
        ['"自认中国人"比例', '2.4%（1992年以来新低）[4]', '反向同温层（中国认同者）被边缘化至个位数',
         '政大选研'],
        ['"既是台湾人也是中国人"', '约 31%（2024/12）[5]', '中间层稳定但话语权下降', '政大选研'],
        ['"独立支持率"', '44.3%（2025/11，创赖清德上任后新低）[6]',
         '同温层内部出现温和分化，但未跨界', '台湾民意基金会'],
        ['"永远维持现状"', '33.2%（1994年以来新高）[5]',
         '反映同温层对赖政府激进路线的隐性抗拒', '政大选研'],
    ], highlight_col=1)

add_h('1.2 政党同温层的结构性切割', level=2)
add_p('民意基金会 2025 年 12 月调查显示，民进党、国民党、民众党三大政党认同的"叠加重合度"'
      '降至历史低位，三层同温层几乎不交集 [7]：', align='justify')
add_p('· 绿营同温层：约 32–35%（DPP铁盘 + 倾绿独立选民）[7]，主导平面媒体（自由时报、三立、'
      '民视）+ 内容农场转型的 YouTube 网红群（如八卦版、podcast 政治节目）[8]；', align='justify')
add_p('· 蓝营同温层：约 22–28%，主导 TVBS、中天、中时电子报、联合报，以及部分长青粉专（如赵少康频道）；'
      '内部又分为本土蓝（侯友宜系）与深蓝（韩国瑜系），同温层之间也有断层；', align='justify')
add_p('· 白营/民众党同温层：约 12–18%，以柯文哲粉丝群体为核心，依托 PTT、Dcard、'
      'TikTok 青年用户 [8]，是绿蓝两大同温层最难"破圈"的对象；', align='justify')
add_p('· "无政党认同"群体：仍占 30% 左右——但 2025 年大罢免投票显示，'
      '这一群体在动员关键时刻向蓝白靠拢，间接证伪了民进党"沉默多数支持自己"的预设 [9][10]。',
      align='justify')

add_h('1.3 演算法驱动的同温层加深机制', level=2)
add_p('RFA 2025 年 4 月深度报道揭示，中共统战部"会分析特定网红的后台数据，对接单的网红给予'
      '具体指令制作特定议题影片"——多位原本讨论经济议题的网红，'
      '在印度移工议题上突然使用"一模一样的台词" [8]。这一发现的反向意义是：', align='justify')
add_p('① 同温层不是自然形成的，而是被算法+议题设定（agenda setting）双重塑造；',
      align='justify')
add_p('② TikTok / YouTube Shorts / Threads 的演算法在 2024 年大选后成为'
      '"破圈"或"固圈"的关键基础设施 [8]；', align='justify')
add_p('③ 对绿营而言，演算法是"同温层加固器"——只要保住己方同温层动员效率，'
      '即便民意基础不到 35%，也可在选举/罢免等关键节点维持治理。', align='justify')

# ============ II 赖清德两岸认知作战事件谱 ============
add_h('二、赖清德两岸认知作战事件谱（2024.05–2026.05）', level=1)
add_p('以下事件按"政治-法律-文化-教育-传播"五条战线展开，构成一套完整的认知重塑工程。',
      align='justify')

add_h('2.1 政治-法律战线：从"互不隶属"到"境外敌对势力"', level=2)
add_table(
    ['日期', '事件', '关键定性'],
    [
        ['2024.05.20', '赖清德就职演说提出"中华民国与中华人民共和国互不隶属"，被陆方定性为"新两国论" [11][12]',
         '执政基调的"独化"明文化'],
        ['2024.10.10', '"双十讲话"重申两岸互不隶属，并妄称"中华人民共和国无权代表台湾" [13][14]',
         '将"两国论"升级为"主权排他论"'],
        ['2025.03.13', '高层国安会议后公布"17项策略"，首次将大陆正式定义为"境外敌对势力"，'
         '恢复军事审判制度防渗透 [15][16][17]', '法律层面的"敌国化"动作'],
        ['2025.05–08', '"大罢免"两波启动，目标 31 位国民党立委 + 民众党新竹市长，'
         '动员公民团体配合 [9][10][18]', '将认知作战延伸至选举工具'],
        ['2025.08.23', '"大罢免"全部失败，DPP 投入数十亿新台币、半年时间、行政与司法手段'
         '全部告负 [9][10]', '同温层无法破圈的实证'],
    ])

add_h('2.2 文化战线：从"重建台湾艺术史"到中华文化总会更名', level=2)
add_table(
    ['日期', '事件', '内容'],
    [
        ['2024.09', '台文化主管部门抛出"重建台湾艺术史 2.0 计划"，2025–2028 共投入 14 亿元新台币，'
         '聚焦美术史、音乐史、文学史、建筑史、歌仔戏史五大类别 [19][20]',
         '将"去中国化"延伸至艺术领域'],
        ['2025', '台文化主管部门年度预算 290.09 亿元，加前瞻基建 13.84 亿 = 303.93 亿元，'
         '首度突破 300 亿，年增近 12% [19]', '财政上对"文化台独"的大规模加持'],
        ['2024.05.20', '赖清德兼任"中华文化总会"第八任会长 [21][22]',
         '行政首长直接掌控对外文化窗口'],
        ['2026.03.17', '"中华文化总会"完成更名程序，英文名改为 National Cultural Association of '
         'Taiwan（NCAT）[21]', '组织层面"去中华化"的标志性动作'],
        ['2024–', '"国家文化记忆库"持续扩大，强调"台湾主体性"叙事，'
         '部分博物馆与地方文化馆把"中华文化根脉"叙事弱化为"多元来源之一" [23]',
         '叙事层面"主体性"对"根脉性"的替代'],
    ])

add_blockquote('观察：文化总会的更名不是终点，而是"系统工程"的标志节点——'
               '将 1967 年蒋介石创立"中华文化复兴运动推行委员会"以来沿用近 60 年的'
               '"中华"招牌正式撤下，相当于对"两岸同根同源"叙事的制度性切断。')

add_h('2.3 语言战线：以"台语 / 台湾客语"挤压"国语"', level=2)
add_p('· 2019 年《国家语言发展法》施行，将"台湾台语""台湾客语""台湾原住民族语"'
      '"马祖语""台湾手语"列为国家语言 [24][25]；', align='justify')
add_p('· 2022 年 7 月行政院核定"国家语言整体发展方案"，要求公文书面建议用语优先使用上述本土语言 [26]；',
      align='justify')
add_p('· 108 课纲（2019 年实施）下，国小阶段本土语文/新住民语文每周 1 节（vs 国语 5–6 节）；'
      '国中/高中阶段降为选修 [25]；', align='justify')
add_p('· 2024 年文化部进一步推动"台语"作为正式行政、教育与公共媒体用语，'
      '将"国语"从制度上"去中心化"——表面是"多元",实质是把'
      '"汉语普通话"压缩为多种语言中的一种 [24][25][26]；', align='justify')
add_p('· 实际语言使用现状（语料库统计）[25]：1986–1994 年出生人口仅 22.3% 以台语为主要语言；'
      '6–14 岁人群仅 7.4%——政策导向与自然语言演变方向相反，'
      '凸显其"政治建构"而非"语言保护"的本质。', align='justify')

add_blockquote('反讽点：当"台语"被升格为国家语言之时，台湾年轻一代台语使用率已降至 7.4%。'
               '这意味着"语言台独"工程的真实功能不是"复振台语"，'
               '而是借由制度操作弱化普通话/汉语在公共领域的中心地位。')

add_h('2.4 教育战线：108 课纲与"识读中国"教材', level=2)
add_p('· 108 课纲 2019 年实施。高中历史教科书原"中国史"章节直接消失，被并入"东亚史"，'
      '与日韩并列 [27][28][29]；', align='justify')
add_p('· 小学历史无"中国史"专章；初中 18 个单元中涉中国仅 3 个，占 1/6 [27][29]；',
      align='justify')
add_p('· 高中文言文比例从 45–55% 调降至 35–45%，删除荀子《劝学》、屈原《渔父》、'
      '范仲淹《岳阳楼记》、欧阳修《醉翁亭记》、顾炎武《廉耻》等 17 篇经典 [27][28]；',
      align='justify')
add_p('· 2025 年 7 月，台教育部门推出所谓"识读中国"教材，被陆方批评为'
      '"文化台独又一邪恶招数" [30][31]；', align='justify')
add_p('· 国民党立委韩国瑜（立法院长）多次提案恢复"中华文化基本教材"为高中必修，'
      '在立法院教委会与教育部门拉锯 [27][28][30]。', align='justify')

add_h('2.5 传播-管控战线：对赴陆艺人与跨境传播的政治化', level=2)
add_table(
    ['日期', '事件', '处理依据'],
    [
        ['2024.10', '欧阳娜娜等台湾艺人在"联合利剑-2024B"军演期间转发支持军演贴文 [32]',
         '陆方主动议题设定'],
        ['2025.03.11', '陆委会查处转发"台湾必归"贴图的艺人，调查是否与中国党政军合作 [33]',
         '《两岸人民关系条例》第 33 条之一'],
        ['2025.05.15–16', '陆委会主委邱垂正宣布锁定 20 余位"亲陆艺人"调查，包含欧阳娜娜等 [34][35][36]',
         '最高罚锾 10–50 万新台币，可连续处罚'],
        ['2025.08.26', '陆委会公开呼吁台湾艺人勿呼应中共"九三阅兵"统战宣传 [37]',
         '从"事后处罚"扩展到"事前威慑"'],
        ['2025.04', 'RFA 系列报道揭露"中共外宣转型 YouTube/TikTok 影音网红"机制 [8][38]',
         '两岸认知战进入"演算法+网红"阶段'],
    ])

add_blockquote('结构性变化：赖政府对艺人言行的管控，本质上是"对己方同温层之外'
               '具有破圈能力的高曝光个人"的精准压制。被列入名单的不是反对赖清德的人，'
               '而是有能力跨同温层向年轻群体传播大陆好感的人。这是"同温层维护工程"的'
               '关键拼图——堵住破圈通道。')

# ============ III 议题破圈规律 ============
add_h('三、议题破圈规律：什么样的议题能跨越同温层？', level=1)

add_h('3.1 破圈失败案例 vs 成功案例', level=2)
add_table(
    ['议题类型', '案例', '破圈结果', '原因'],
    [
        ['党派攻防型', '大罢免（2025/7–8）', '失败 [9][10]',
         '议题被定义为"亲共 vs 反共"，反而强化各阵营同温层'],
        ['制度变更型', '17 项策略（2025/3）[15][16]', '低破圈',
         '法律技术性强，绿营内部认同，但中间选民冷淡'],
        ['身份认同型', '"互不隶属"（2024/5）[11]', '中度破圈但反弹强',
         '触及统独红线，但"永远维持现状"反向上升至 33.2% [5]'],
        ['文化符号型', '中华文化总会更名（2026/3）[21]', '低破圈',
         '机构性更名，普通民众感知度低，但精英层强烈对立'],
        ['民生分配型', '中南部无薪假/AI 经济极化（参见经济卷）',
         '高破圈潜力但被绿营议题压制', '经济议题天然跨同温层，但缺少强势主导者'],
        ['国安恐惧型', '艺人调查/赴陆管控 [34]', '中度破圈',
         '激活蓝营反弹同时强化绿营"国家安全"叙事'],
    ])

add_h('3.2 破圈的三个学理条件', level=2)
add_p('A. 议题"切身性"高于"立场性"——民生议题（无薪假、房价、医疗）天然跨越政党同温层；'
      '统独议题则反而加固同温层；', align='justify')
add_p('B. 议题"载体跨平台"——能同时出现在 LINE 群组、TikTok、YouTube、PTT、Threads、'
      '电视新闻六个以上平台的议题，破圈概率显著高于单一平台议题 [8][38]；', align='justify')
add_p('C. 议题"情感模糊性"——纯粹立场议题（蓝绿对错）锁死同温层；'
      '带有道德两难、人情压力、世代焦虑的议题更易破圈（如"年轻人没钱结婚生小孩"）。', align='justify')

add_h('3.3 赖清德团队的议题设定策略', level=2)
add_p('从 2024.05 至 2026.05 的事件密度看，赖政府呈现出以下议题节奏规律：', align='justify')
add_p('· 高敏感期（选举/罢免/国际节点前）：密集投放"国安/敌对势力/认知战"议题，'
      '激活己方同温层动员；', align='justify')
add_p('· 低敏感期：转向文化/语言/教育长效议题，做"温水煮青蛙"式的叙事重构；', align='justify')
add_p('· 危机回避期：对中南部无薪假、新台币升值伤害传统产业、台积电赴美等议题'
      '（参见经济卷）做"低声化"处理，避免出现可破圈的民生议题。', align='justify')

add_blockquote('政治启示：赖政府对议题的精准切换显示其团队深谙同温层政治的核心——'
               '不是赢得多数，而是确保己方多数动员、对方多数沉默。'
               '同温层固化在此被当作治理资源，而非政治病症。')

# ============ IV 双重撕裂 ============
add_h('四、同温层固化与经济 K 型撕裂的"双重撕裂"耦合', level=1)
add_p('本卷与经济卷的核心交叉发现：台湾正在经历同步的"经济撕裂"+"舆论撕裂"。'
      '两者并非平行，而是相互强化：', align='justify')

add_h('4.1 经济同温层 ↔ 政治同温层的重合', level=2)
add_table(
    ['维度', '上层（赢家）', '下层（输家）'],
    [
        ['经济卷', '台积电基层工程师年薪 200–300 万新台币；新竹/台北科技走廊；电子业薪资'
         '为全经济均值 1.7 倍（参见经济卷 [5][6]）',
         '中南部传统制造业工人年薪 50–60 万；2025 年底 8,331 人实施无薪假'
         '（95% 集中在制造业）（参见经济卷 [42][43]）'],
        ['政治认同', '深绿都会专业阶层；高度认同"台湾主体性"叙事；'
         '主流接收信息为自由时报、三立、绿营 YouTuber',
         '蓝营/中间偏蓝、本土传统业；对赖政府"国安叙事"反应冷淡；'
         '关心民生议题但议题被压制'],
        ['文化资本', '掌握英语/赴美机会/海外 fab 派任；与硅谷/美方决策圈对接',
         '本土汉语+台语双语；与大陆既有商贸/亲属/文化网络断裂代价高'],
        ['同温层维持成本', '低（主流媒体+演算法天然倾斜）', '高（被边缘化、'
         '主流叙事中被定性为"被统战""旧时代"）'],
    ])

add_h('4.2 "双重撕裂"的政治后果', level=2)
add_p('· 当经济上的上下层与政治认同上的绿/蓝白层高度重合，社会无法通过"经济议题动员中间选民"'
      '来打破政党同温层——这就是大罢免失败的深层原因 [9][10]；', align='justify')
add_p('· 经济撕裂为政治同温层提供了物质基础：赢家天然支持"维持现状+全面对美绑定"，'
      '输家天然怀疑这一路线——但输家的怀疑找不到强势政治载体；', align='justify')
add_p('· 一旦 AI Capex 周期出现回落（参见经济卷第 7.3 节），'
      '"经济输家+政治输家"的双重叠加群体可能突然扩大，形成同温层结构的'
      '"压力性破裂"——届时议题破圈方向可能从"国安议题"转向"民生分配议题"，'
      '执政党对议题的主导权会显著下降。', align='justify')

# ============ V 趋势 ============
add_h('五、2026–2028 趋势研判', level=1)

add_h('5.1 同温层固化的三重持续动力', level=2)
add_p('① 制度性动力：17 项策略、军事审判恢复、艺人调查 → 行政手段持续清理破圈通道 [15][34]；',
      align='justify')
add_p('② 演算法动力：TikTok/YouTube/Threads 的推送机制将持续强化同温层 [8][38]；', align='justify')
add_p('③ 文化叙事动力：重建艺术史、文化总会更名、教材去中国化等长效工程进入收获期 [19][21][27]。',
      align='justify')

add_h('5.2 同温层固化的三重潜在破口', level=2)
add_p('① 经济周期破口：AI Capex 2027–2028 若回落，无薪假/中小企业问题外溢成主流议题；',
      align='justify')
add_p('② 世代破口：Z 世代（2010 年后出生）在 2028 年大选成为新选民，'
      '其同温层不在传统蓝绿地图上，民众党/新政治力量的崛起空间；', align='justify')
add_p('③ 国际破口：若美方"50:50 半导体分配"加速、特朗普 2.0 的不可预测性放大，'
      '台湾对美完全依附路线在岛内可能首次出现质疑——这是赖政府最难掌控的破口。',
      align='justify')

# ============ VI 结论 ============
add_h('六、结论', level=1)
add_p('赖清德上任以来的台湾岛内政治-文化生态，可概括为"同温层固化+议题精准切换+破圈通道清理"'
      '的复合体系。其本质是一场"以同温层为治理工具"的政治工程：', align='justify')

add_p('第一层（同温层层）', bold=True)
add_p('"自认台湾人 61.7%、自认中国人 2.4%" 的认同分布已被叙事工程固化为"主流共识"，'
      '反向同温层被边缘化至个位数；同时绿/蓝/白三大政党同温层重合度持续下降 [4][7]。',
      align='justify')

add_p('第二层（事件谱层）', bold=True)
add_p('从"互不隶属"（2024/5）到"境外敌对势力"（2025/3）到"17 项策略"再到艺人调查（2025/5）、'
      '中华文化总会更名（2026/3），构成一条政治-法律-文化-教育-传播五线并进的'
      '"系统性认知重塑工程" [11][15][21][34]。', align='justify')

add_p('第三层（双重撕裂层）', bold=True)
add_p('文化卷的同温层固化与经济卷的 K 型撕裂高度耦合，构成"双重撕裂"。'
      '同温层固化使经济输家无法政治化，经济极化又为同温层提供物质基础——'
      '形成自我强化的政治-经济双螺旋。', align='justify')

add_p('最终研判', bold=True)
add_p('赖政府的同温层治理策略在 2024–2026 取得了短期成功（巩固铁盘 + 阻止破圈 + 制度化敌意），'
      '但同时也付出了三重代价：①"永远维持现状"民意创新高（33.2%），反映温和民意对激进路线的隐性抗拒；'
      '② 大罢免完败，证明同温层无法外扩；③ 经济输家群体被结构性排除，'
      '埋下未来"民生议题反向破圈"的火种。2027–2028 一旦 AI 周期下行 + 美方供应链再平衡 + '
      '世代更替三重压力同时叠加，当前的同温层固化体系将面临系统性应力测试。',
      align='justify')

# ============ References ============
add_h('参考文献（References, APA-modified）', level=1)
refs = [
    "[1] Pariser, E. (2011). The Filter Bubble: What the Internet Is Hiding from You. "
    "Penguin Press. （概念引自维基百科“过滤气泡”条目）https://zh.wikipedia.org/wiki/%E9%81%8E%E6%BF%BE%E6%B0%A3%E6%B3%A1",
    "[2] Sunstein, C. R. (2001, 2006). Republic.com & Infotopia. "
    "（“回声室效应”与“信息茧房”概念起源）https://zh.wikipedia.org/zh-tw/%E8%BF%B4%E8%81%B2%E5%AE%A4%E6%95%88%E6%87%89",
    "[3] Fondation Descartes. (2020). Filter bubbles and echo chambers. "
    "https://www.fondationdescartes.org/en/2020/07/filter-bubbles-and-echo-chambers/",
    "[4] 中央社（CNA）. (2024年2月23日). 政大調查：民眾自認中國人比率創新低 "
    "自認台灣人近4年都逾6成. https://www.cna.com.tw/news/aipl/202402230063.aspx",
    "[5] 國立政治大學選舉研究中心. (2024–2025). 臺灣民眾臺灣人/中國人認同趨勢分佈. "
    "https://esc.nccu.edu.tw/PageDoc?fid=7804",
    "[6] Newtalk新聞. (2025年11月13日). 台灣民意基金會民調：支持台獨降至44.3% "
    "創賴清德上任後新低. https://newtalk.tw/news/view/2025-11-13/1004238",
    "[7] 台灣民意基金會. (2025年12月23日). 2025年12月全國性民意調查摘要報告：“台灣人眼中的民主進步黨”"
    "（游盈隆）. https://www.tpof.org/wp-content/uploads/2025/12/",
    "[8] RFA中文. (2025年4月21日). 深度报道｜中共外宣在台湾之九：内容农场转型影音网红. "
    "https://www.rfa.org/mandarin/shishi-hecha/2025/04/21/fact-check-ccp-propaganda-funseetv/",
    "[9] 任冬梅. (2025). 民进党“大罢免，大失败”的原因. 中国社科院台湾研究所. "
    "http://its.taiwan.cssn.cn/cgzs/zyhd/202509/t20250924_5916890.shtml",
    "[10] RFA中文. (2025年7月23日). 台湾26日投票“大罢免”：中国因素推波助澜？ "
    "https://www.rfa.org/mandarin/guoji/yatai/gangtai/2025/07/23/",
    "[11] 维基百科. (2024). 2024年中华民国总统就职典礼. "
    "https://zh.wikipedia.org/zh-hans/2024%E5%B9%B4%E4%B8%AD%E8%8F%AF%E6%B0%91%E5%9C%8B%E7%B8%BD%E7%B5%B1%E5%B0%B1%E8%81%B7%E5%85%B8%E7%A6%AE",
    "[12] 光明日报. (2024年5月23日). 赖清德“就职演说”充斥着赤裸裸的“台独”挑衅. "
    "https://news.gmw.cn/2024-05/23/content_37337954.htm",
    "[13] 新华网. (2024年10月11日). “新两国论”的再包装——台湾各界批驳赖清德“双十”讲话. "
    "http://www.news.cn/20241011/24aeaf0a6dac4f448afcd937a2585694/c.html",
    "[14] 求是网. (2024年10月14日). 赖清德“新两国论”严重危害台海和平. "
    "https://www.qstheory.cn/qshyjx/2024-10/14/c_1130209542.htm",
    "[15] RFA中文. (2025年3月13日). 赖清德称中国是“境外敌对势力” 恢复军事审判制度防渗透. "
    "https://www.rfa.org/mandarin/guoji/yatai/gangtai/2025/03/13/laichingte-taiwan-militarytrial-prevent-ccp-infiltration/",
    "[16] 经济观察网. (2025年3月17日). 台海观澜｜赖清德当局出台“拒统17条”，但不敢触碰红线. "
    "http://www.eeo.com.cn/2025/0317/716874.shtml",
    "[17] 国台办. (2025年3月26日). 赖清德抛出“17项策略”彻底暴露“绿色恐怖”丑恶面目. "
    "http://www.vos.com.cn/haixiah5/2025-03/26/cms219008article.shtml",
    "[18] 维基百科. (2025). 大罢免（条目）. "
    "https://zh.wikipedia.org/zh-hans/%E5%A4%A7%E7%BD%B7%E5%85%8D",
    "[19] 任冬梅. (2024年11月29日). 赖清德“文化台独工程”正加速推进. 华夏经纬网. "
    "https://www.huaxia.com/c/2024/11/29/1982241.shtml",
    "[20] 新华网. (2023年12月14日). 署名文章：赖清德“文化台独3.0”用心险恶、贻害无穷. "
    "http://www.news.cn/tw/2023-12/14/c_1130026792.htm",
    "[21] 维基百科. (2026). 中华文化总会条目. "
    "https://zh.wikipedia.org/zh-hans/%E4%B8%AD%E8%8F%AF%E6%96%87%E5%8C%96%E7%B8%BD%E6%9C%83",
    "[22] 人民政协网. (2017年3月10日). 蔡英文接掌中华文化总会 被指走向“文化台独”. "
    "https://www.rmzxw.com.cn/c/2017-03-10/1401060.shtml",
    "[23] 中国和平统一促进会. (2024). 台湾“去中国化”教育研究. "
    "http://www.zhongguotongcuhui.org.cn/tylt/202406/202412/t20241230_12674888.html",
    "[24] 台行政院全球资讯网. (2019). 《國家語言發展法》—改善語言斷層危機、尊重多元文化發展. "
    "https://www.ey.gov.tw/Page/5A8A0CB5B41DA11E/acb034c7-e184-4a39-be3f-381db50a6abe",
    "[25] 维基百科. (2025). 国语政策 / 台湾语言保护 条目. "
    "https://zh.wikipedia.org/zh-hans/%E5%9C%8B%E8%AA%9E%E6%94%BF%E7%AD%96",
    "[26] 台文化部. (2022年7月15日). 国家语言整体发展方案（行政院核定）. "
    "https://www.moc.gov.tw/cp.aspx?n=131",
    "[27] 中国和平统一促进会. (2023). 民进党“语言台独”布局及其深层危害. "
    "http://www.zhongguotongcuhui.org.cn/tylt/202304/202311/t20231106_12579523.html",
    "[28] 张方远. (2021). “台独”历史教科书的由来与现状. 紫荆杂志. "
    "https://zijing.com.cn/article/2021-02/08/content_926242877155143680.html",
    "[29] 北京日报客户端. (2024). 一晃5年，“台独课纲”这样贻害岛内青少年. "
    "https://xinwen.bjd.com.cn/content/s665a6f1ae4b035c6ca5e39e8.html",
    "[30] 国台办. (2024年11月). 救救孩子！“台独”课纲犯众怒，民众自救要将其废除. "
    "https://www.gwytb.gov.cn/zt/zylszl/tyqy/202411/t20241128_12667498.htm",
    "[31] iTaiwan News. (2025年7月21日). 再次“洗脑”莘莘学子，所谓“识读中国”教材是民进党当局"
    "推行“文化台独”的又一邪恶招数. https://www.itaiwannews.cn/2025-07-21/",
    "[32] 观察者网. (2025年5月16日). 台当局对艺人下手：已锁定欧阳娜娜等20多个亲大陆艺人. "
    "https://www.guancha.cn/politics/2025_05_16_776037.shtml",
    "[33] RFA中文. (2025年3月11日). 台湾艺人转发“台湾必归” 陆委会将查是否与中国党政军合作. "
    "https://www.rfa.org/mandarin/guoji/yatai/gangtai/2025/03/11/taiwan-china-status-province-independence/",
    "[34] RFA中文. (2025年5月15日). 台湾艺人配合中共发文 陆委会：列20多人 将偕同文化部调查. "
    "https://www.rfa.org/mandarin/xinwenkuaixun/2025/05/15/taiwan-investigate-celebrity-post-about-chinas-united-front/",
    "[35] 百度百科. (2025). 2025年台当局威胁艺人事件. "
    "https://baike.baidu.com/item/2025%E5%B9%B4%E5%8F%B0%E5%BD%93%E5%B1%80%E5%A8%81%E8%83%81%E8%89%BA%E4%BA%BA%E4%BA%8B%E4%BB%B6/65693433",
    "[36] 禁闻网. (2025年5月16日). 台湾艺人配合中共发文 陆委会：将偕同文化部调查20多人. "
    "https://www.bannedbook.org/bnews/ssgc/20250516/2191827.html",
    "[37] 大纪元. (2025年8月26日). 陆委会吁台艺人勿呼应中共九三阅兵统战宣传. "
    "https://www.epochtimes.com/gb/25/8/26/n14581200.htm",
    "[38] RFA中文. (2025年4月28日). 深度报道｜中共外宣在台湾（完结篇）：入岛、入户，入TikTok. "
    "https://www.rfa.org/mandarin/shishi-hecha/2025/04/28/factcheck-ccp-propaganda-tiktok/",
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.first_line_indent = Cm(-0.6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(r)
    run.font.size = Pt(9.5)
    run.font.name = 'Microsoft YaHei'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# Footer
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fr = foot.add_run('编制：文化卷扩展深度研究（v2.0，与经济卷 v2.0 配套）  ·  日期：2026年5月14日')
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
fr.italic = True
fr.font.name = 'Microsoft YaHei'
fr.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.save(OUT)
print('OK:', OUT, 'size=', os.path.getsize(OUT))
