"""Render GetOilFromChina v2.1 situation report as formatted DOCX.

Identical formatting to v2.0 script: FangSong body 14pt, bold headings,
tables 10.5pt, references 9pt with hanging indent, superscript citations.
"""

import os, re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

here = os.path.dirname(os.path.abspath(__file__))
src = os.path.join(here, "07_report_v2.3_situation_report.md")
dst = os.path.join(here, "GetOilFromChina_情况调查报告_v2.3.docx")

with open(src, "r", encoding="utf-8") as f:
    lines = f.read().split("\n")

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

normal = doc.styles["Normal"]
normal.font.name = "FangSong"
normal.font.size = Pt(14)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "FangSong")

CN_FONT = "FangSong"
LATIN_FONT = "Times New Roman"

CITATION_RE = re.compile(r"\[(\d+(?:[,\-–]\s*\d+)*)\]")

def set_cn_font(run, size_pt, bold=False, superscript=False):
    run.font.name = LATIN_FONT
    run.font.size = Pt(size_pt)
    run.bold = bold
    if superscript:
        run.font.superscript = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:eastAsia"), CN_FONT)
    rFonts.set(qn("w:ascii"), LATIN_FONT)
    rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    rFonts.set(qn("w:cs"), LATIN_FONT)

def add_inline_runs(paragraph, text, size, base_bold=False, parse_citations=True):
    bold_split = re.split(r"(\*\*[^*\n]+\*\*)", text)
    for seg in bold_split:
        if not seg:
            continue
        seg_is_bold = seg.startswith("**") and seg.endswith("**") and len(seg) > 4
        seg_text = seg[2:-2] if seg_is_bold else seg
        if parse_citations:
            cite_split = CITATION_RE.split(seg_text)
            for i, part in enumerate(cite_split):
                if i % 2 == 1:
                    run = paragraph.add_run(f"[{part}]")
                    set_cn_font(run, max(9, size - 4), bold=False, superscript=True)
                else:
                    if part:
                        run = paragraph.add_run(part)
                        set_cn_font(run, size, bold=(base_bold or seg_is_bold))
        else:
            run = paragraph.add_run(seg_text)
            set_cn_font(run, size, bold=(base_bold or seg_is_bold))

def add_heading(text, level):
    sizes = {1: 22, 2: 15, 3: 14, 4: 12}
    aligns = {1: WD_ALIGN_PARAGRAPH.CENTER, 2: WD_ALIGN_PARAGRAPH.LEFT,
              3: WD_ALIGN_PARAGRAPH.LEFT, 4: WD_ALIGN_PARAGRAPH.LEFT}
    p = doc.add_paragraph()
    p.alignment = aligns.get(level, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(6)
    add_inline_runs(p, text, sizes.get(level, 14), base_bold=True, parse_citations=False)

def add_body(text, in_references=False):
    p = doc.add_paragraph()
    if in_references:
        p.paragraph_format.left_indent = Pt(28)
        p.paragraph_format.first_line_indent = Pt(-28)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(2)
        add_inline_runs(p, text, 9, base_bold=False, parse_citations=False)
    else:
        p.paragraph_format.first_line_indent = Pt(28)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)
        add_inline_runs(p, text, 14, base_bold=False, parse_citations=True)

def add_bullet(text, in_references=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(28)
    p.paragraph_format.line_spacing = 1.25 if in_references else 1.5
    if in_references:
        add_inline_runs(p, "• " + text, 9, base_bold=False, parse_citations=False)
    else:
        add_inline_runs(p, "• " + text, 14, base_bold=False, parse_citations=True)

def add_blockquote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(28)
    p.paragraph_format.right_indent = Pt(28)
    p.paragraph_format.line_spacing = 1.5
    add_inline_runs(p, text, 14, base_bold=False, parse_citations=True)

def add_table_from_rows(data):
    if not data: return
    cols = max(len(r) for r in data)
    tbl = doc.add_table(rows=len(data), cols=cols)
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = True
    for ri, row in enumerate(data):
        for ci in range(cols):
            cell = tbl.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            cell.text = ""
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing = 1.15
                add_inline_runs(para, text, 10.5, base_bold=(ri == 0), parse_citations=False)
    doc.add_paragraph()

in_refs = False

i = 0
while i < len(lines):
    line = lines[i].rstrip()
    stripped = line.strip()

    # mark references mode (附件六 参考文献 in v2.1)
    if "参考文献" in stripped[:30] or "附件六" in stripped[:30]:
        in_refs = True

    if stripped.startswith("#"):
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            text = m.group(2).strip()
            add_heading(text, level)
            i += 1
            continue

    if stripped == "---":
        doc.add_paragraph()
        i += 1
        continue

    if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
        rows = []
        while i < len(lines) and lines[i].lstrip().startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            rows.append(cells)
            i += 1
        if len(rows) >= 2:
            data = [rows[0]] + rows[2:]
            add_table_from_rows(data)
        continue

    if stripped.startswith("- "):
        add_bullet(stripped[2:].strip(), in_references=in_refs)
        i += 1
        continue

    nm = re.match(r"^\d+\.\s+(.*)$", stripped)
    if nm:
        add_bullet(nm.group(1), in_references=in_refs)
        i += 1
        continue

    if stripped.startswith(">"):
        add_blockquote(stripped.lstrip("> ").strip())
        i += 1
        continue

    if stripped:
        add_body(stripped, in_references=in_refs)
    i += 1

doc.save(dst)
print(f"Saved: {dst}  ({os.path.getsize(dst)} bytes)")
