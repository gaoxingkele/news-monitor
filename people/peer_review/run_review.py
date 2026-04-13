"""三专家评审主入口。

用法：
    python -m people.peer_review.run_review \
        --report people/蒋万安_政治人物分析报告_v1.0.docx \
        --person 蒋万安 \
        --output people/peer_review/reports/蒋万安_v1.0_peer_review.md
"""
from __future__ import annotations

import argparse
import datetime as dt
import json
import logging
import sys
from pathlib import Path

# 允许从项目根直接运行
sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from people.peer_review.config import build_channels, bootstrap_env
from people.peer_review.llm_clients import LLMCallError, call_expert
from people.peer_review.prompts import EXPERTS, build_user_prompt


logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("peer_review.run")


def extract_docx_text(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    # 简单把表格内容也带上
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def run_reviews(
    report_path: Path,
    person: str,
    output_path: Path,
    only: set[str] | None = None,
) -> None:
    bootstrap_env()

    raw_dir = Path(__file__).resolve().parent / "raw_responses"
    raw_dir.mkdir(parents=True, exist_ok=True)

    if report_path.suffix.lower() == ".docx":
        report_text = extract_docx_text(report_path)
    else:
        report_text = report_path.read_text(encoding="utf-8")

    logger.info("loaded report: %s (chars=%d)", report_path, len(report_text))

    channels = build_channels()

    results: dict[str, dict] = {}
    # 若设置了 only，未选中的专家从已有 raw_responses 中加载，避免重复调用
    for ch in channels:
        if only and ch.expert_id not in only:
            report_tag = report_path.stem.replace(" ", "_")
            raw_file = raw_dir / f"{report_tag}_expert_{ch.expert_id}_{ch.provider}.md"
            if raw_file.is_file():
                text = raw_file.read_text(encoding="utf-8")
                # 去掉首行标题
                lines = text.splitlines()
                if lines and lines[0].startswith("#"):
                    text = "\n".join(lines[1:]).strip()
                results[ch.expert_id] = {
                    "display_name": ch.display_name,
                    "channel": ch.provider,
                    "model": ch.model,
                    "content": text,
                    "meta": {"loaded_from_cache": True},
                    "error": None,
                }
                logger.info("[%s] loaded from cache: %s", ch.expert_id, raw_file.name)
            continue
        system_prompt = EXPERTS[ch.expert_id]
        user_prompt = build_user_prompt(person, report_text, expert_id=ch.expert_id)
        logger.info("=== 调用 %s (%s) ===", ch.display_name, ch.model)
        try:
            content, meta = call_expert(ch, system_prompt, user_prompt)
            results[ch.expert_id] = {
                "display_name": ch.display_name,
                "channel": ch.provider,
                "model": ch.model,
                "content": content,
                "meta": meta,
                "error": None,
            }
            # 保存原始响应
            report_tag = report_path.stem.replace(" ", "_")
            raw_file = raw_dir / f"{report_tag}_expert_{ch.expert_id}_{ch.provider}.md"
            raw_file.write_text(
                f"# {ch.display_name} ({ch.provider} / {ch.model})\n\n{content}\n",
                encoding="utf-8",
            )
        except LLMCallError as e:
            logger.error("[%s] 调用失败: %s", ch.expert_id, e)
            results[ch.expert_id] = {
                "display_name": ch.display_name,
                "channel": ch.provider,
                "model": ch.model,
                "content": "",
                "meta": {},
                "error": str(e),
            }

    write_aggregate_report(output_path, person, report_path, results)
    logger.info("评审报告已输出：%s", output_path)

    # 以简要 json meta 打印摘要
    summary = {
        k: {"model": v["model"], "channel": v["channel"], "error": v["error"],
            "len": len(v["content"])}
        for k, v in results.items()
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))


def write_aggregate_report(
    output_path: Path,
    person: str,
    report_path: Path,
    results: dict[str, dict],
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    now = dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    header = f"""# {person} 政治人物分析报告 v1.0 · 三专家评审

- **被评审报告**: `{report_path.name}`
- **评审框架基线**: OSINT V6.0 全景分析框架
- **评审时间**: {now}
- **评审组成**:
  - 国际关系专家A → {results.get("A", {}).get("channel", "-")} / `{results.get("A", {}).get("model", "-")}`
  - 台湾历史学家B → {results.get("B", {}).get("channel", "-")} / `{results.get("B", {}).get("model", "-")}`
  - 台湾问题研究专家C → {results.get("C", {}).get("channel", "-")} / `{results.get("C", {}).get("model", "-")}`
  - 情报收集专家D（凡博士）→ {results.get("D", {}).get("channel", "-")} / `{results.get("D", {}).get("model", "-")}`

---
"""

    sections: list[str] = [header]
    for key in ("A", "B", "C", "D"):
        r = results.get(key)
        if not r:
            continue
        title_line = f"# 评审意见 · {r['display_name']}\n"
        meta_line = (
            f"> 通道: {r['channel']} · 模型: `{r['model']}` · "
            f"tokens: {r['meta'].get('usage', {})}\n"
        )
        if r["error"]:
            body = f"\n> **调用失败**: {r['error']}\n"
        else:
            body = "\n" + r["content"].strip() + "\n"
        sections.append(title_line + meta_line + body + "\n---\n")

    output_path.write_text("\n".join(sections), encoding="utf-8")


def main() -> None:
    ap = argparse.ArgumentParser(description="OSINT V6 三专家评审流水线")
    ap.add_argument(
        "--report",
        required=True,
        help="待评审报告路径（.docx 或 .md/.txt）",
    )
    ap.add_argument("--person", required=True, help="评审对象人物姓名")
    ap.add_argument(
        "--output",
        required=True,
        help="汇总评审 markdown 输出路径",
    )
    ap.add_argument(
        "--only",
        default="",
        help="仅执行指定专家 ID，逗号分隔，如 A,C；未选中的专家从 raw_responses 缓存回填",
    )
    args = ap.parse_args()

    only_set: set[str] | None = None
    if args.only.strip():
        only_set = {x.strip().upper() for x in args.only.split(",") if x.strip()}

    run_reviews(
        report_path=Path(args.report).resolve(),
        person=args.person,
        output_path=Path(args.output).resolve(),
        only=only_set,
    )


if __name__ == "__main__":
    main()
