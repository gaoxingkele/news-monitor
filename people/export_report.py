"""将 markdown 报告导出为 Word (.docx) 和 PDF。

依赖：python-docx, reportlab
用法：python people/export_report.py --input people/output/蒋万安/v1.6.md
"""
import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak,
    Table, TableStyle, KeepTogether,
)
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


# ────────────────── Markdown 解析 ──────────────────

def parse_md(text: str):
    """简易 Markdown 解析，返回 (type, content) 列表。"""
    blocks = []
    in_table = False
    table_rows = []

    for line in text.split("\n"):
        stripped = line.strip()

        # 表格行
        if stripped.startswith("|") and "|" in stripped[1:]:
            # 跳过分隔行
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                continue
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(cells)
            continue
        else:
            if in_table:
                blocks.append(("table", table_rows))
                in_table = False
                table_rows = []

        if not stripped:
            continue

        # 标题
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            blocks.append((f"h{level}", m.group(2).strip("*").strip()))
            continue

        # 分隔线
        if re.match(r"^-{3,}$", stripped):
            blocks.append(("hr", ""))
            continue

        # 粗体标题行（**第X篇**）
        m2 = re.match(r"^\*\*(第[一二三四五六七八]篇.*?)\*\*$", stripped)
        if m2:
            blocks.append(("h2", m2.group(1)))
            continue

        # 普通段落
        blocks.append(("para", stripped))

    if in_table:
        blocks.append(("table", table_rows))

    return blocks


# ────────────────── Word 导出 ──────────────────

def _clean(text: str) -> str:
    """去除 markdown 格式符号。"""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    return text


def export_docx(blocks, out_path: Path):
    doc = Document()

    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # 样式
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "宋体"
    style_normal.font.size = Pt(10.5)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)

    for block_type, content in blocks:
        if block_type == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(_clean(content))
            run.bold = True
            run.font.size = Pt(22)
            run.font.name = "黑体"
        elif block_type == "h2":
            p = doc.add_paragraph()
            run = p.add_run(_clean(content))
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = "黑体"
            p.paragraph_format.space_before = Pt(18)
        elif block_type == "h3":
            p = doc.add_paragraph()
            run = p.add_run(_clean(content))
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = "黑体"
            p.paragraph_format.space_before = Pt(12)
        elif block_type == "h4":
            p = doc.add_paragraph()
            run = p.add_run(_clean(content))
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(8)
        elif block_type == "para":
            p = doc.add_paragraph(_clean(content))
        elif block_type == "table":
            if not content:
                continue
            ncols = max(len(row) for row in content)
            table = doc.add_table(rows=len(content), cols=ncols)
            table.style = "Light Grid Accent 1"
            for i, row in enumerate(content):
                for j, cell in enumerate(row):
                    if j < ncols:
                        table.rows[i].cells[j].text = _clean(cell)
                        for para in table.rows[i].cells[j].paragraphs:
                            para.style.font.size = Pt(9)
        elif block_type == "hr":
            pass  # skip

    doc.save(str(out_path))
    print(f"DOCX → {out_path} ({out_path.stat().st_size // 1024} KB)")


# ────────────────── PDF 导出 ──────────────────

def _find_chinese_font():
    """查找系统中可用的中文字体。"""
    candidates = [
        ("SimSun", r"C:\Windows\Fonts\simsun.ttc"),
        ("SimHei", r"C:\Windows\Fonts\simhei.ttf"),
        ("MicrosoftYaHei", r"C:\Windows\Fonts\msyh.ttc"),
        ("NSimSun", r"C:\Windows\Fonts\nsimsun.ttf"),
    ]
    for name, path in candidates:
        if Path(path).exists():
            try:
                pdfmetrics.registerFont(TTFont(name, path))
                return name
            except Exception:
                continue
    return "Helvetica"


def export_pdf(blocks, out_path: Path):
    font_name = _find_chinese_font()
    print(f"PDF font: {font_name}")

    styles = getSampleStyleSheet()

    style_title = ParagraphStyle(
        "CNTitle", parent=styles["Title"],
        fontName=font_name, fontSize=20, alignment=TA_CENTER,
        spaceAfter=20,
    )
    style_h2 = ParagraphStyle(
        "CNH2", parent=styles["Heading2"],
        fontName=font_name, fontSize=15, spaceBefore=20, spaceAfter=10,
        textColor=colors.HexColor("#1a3c6e"),
    )
    style_h3 = ParagraphStyle(
        "CNH3", parent=styles["Heading3"],
        fontName=font_name, fontSize=12, spaceBefore=14, spaceAfter=6,
        textColor=colors.HexColor("#2c5f8a"),
    )
    style_body = ParagraphStyle(
        "CNBody", parent=styles["Normal"],
        fontName=font_name, fontSize=10, leading=16,
        alignment=TA_JUSTIFY, spaceAfter=6,
        firstLineIndent=20,
    )
    style_meta = ParagraphStyle(
        "CNMeta", parent=styles["Normal"],
        fontName=font_name, fontSize=9, leading=14,
        textColor=colors.grey,
    )
    style_table_cell = ParagraphStyle(
        "CNCell", parent=styles["Normal"],
        fontName=font_name, fontSize=8, leading=11,
    )

    doc = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
    )

    story = []

    for block_type, content in blocks:
        clean = _clean(content) if isinstance(content, str) else content

        if block_type == "h1":
            story.append(Paragraph(clean, style_title))
        elif block_type == "h2":
            story.append(Paragraph(clean, style_h2))
        elif block_type == "h3":
            story.append(Paragraph(clean, style_h3))
        elif block_type == "h4":
            story.append(Paragraph(f"<b>{clean}</b>", style_body))
        elif block_type == "para":
            # 元信息行
            if clean.startswith("**") and "：" in clean[:20]:
                story.append(Paragraph(_clean(clean), style_meta))
            else:
                story.append(Paragraph(clean, style_body))
        elif block_type == "table":
            if not content:
                continue
            ncols = max(len(r) for r in content)
            # 构建表格数据
            tdata = []
            for row in content:
                cells = [Paragraph(_clean(c), style_table_cell) for c in row]
                while len(cells) < ncols:
                    cells.append(Paragraph("", style_table_cell))
                tdata.append(cells)

            avail_width = A4[0] - 5 * cm
            col_w = avail_width / ncols

            t = Table(tdata, colWidths=[col_w] * ncols)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d6e4f0")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1a3c6e")),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(KeepTogether([t, Spacer(1, 8)]))
        elif block_type == "hr":
            story.append(Spacer(1, 12))

    doc.build(story)
    print(f"PDF → {out_path} ({out_path.stat().st_size // 1024} KB)")


# ────────────────── Main ──────────────────

if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="导出 MD 报告为 DOCX + PDF")
    ap.add_argument("--input", required=True, help="输入 markdown 文件路径")
    ap.add_argument("--output-dir", default=None, help="输出目录（默认与输入文件同目录）")
    args = ap.parse_args()

    src = Path(args.input)
    out_dir = Path(args.output_dir) if args.output_dir else src.parent
    stem = src.stem

    md_text = src.read_text(encoding="utf-8")
    blocks = parse_md(md_text)
    print(f"Parsed {len(blocks)} blocks from {src.name}")

    export_docx(blocks, out_dir / f"{stem}.docx")
    export_pdf(blocks, out_dir / f"{stem}.pdf")

    print("\nDone!")
