"""
蒋万安 OSINT 政治人物分析报告 v1.8 → Word (.docx)
中文期刊论文格式：仿宋GB2312 四号 / 1.2行距 / 首行缩进 / 标题层级醒目
"""
import re, os, sys, argparse
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC    = r"D:\aicoding\news-monitor\people\output\李四川\李四川_政治人物分析报告_v1.8.md"
OUT_DIR = r"D:\aicoding\news-monitor\people\output\李四川"
OUT_DOCX = None  # overrideable via CLI

# ── 中文学术期刊格式常量 ──────────────────────────────
BODY_FONT   = "仿宋"          # 正文：仿宋GB2312
HEAD_FONT   = "黑体"          # 标题：黑体
BODY_SIZE   = Pt(14)          # 四号 ≈ 14pt
HEAD1_SIZE  = Pt(18)          # 一级标题 18pt
HEAD2_SIZE  = Pt(16)          # 二级标题 16pt
HEAD3_SIZE  = Pt(14)          # 三级标题 14pt（加粗）
CAPTION_SIZE= Pt(12)          # 表格标题
LINE_SPACING = 1.2            # 1.2倍行距
INDENT_EM   = Cm(0.74)        # 首行缩进2字符≈0.74cm

# ── 字体设置辅助 ─────────────────────────────────────
def set_run_font(run, face, size, bold=False, italic=False, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.name = face
    run.font.size = size
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), face)
    rPr.insert(0, rFonts)

def add_para(doc, text, face, size, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=Pt(0), space_after=Pt(6), indent=None,
             line_spacing=LINE_SPACING, italic=False, color=None,
             keep_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after  = space_after
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = line_spacing
    if indent:
        pf.first_line_indent = indent
    if keep_next:
        pf.keep_with_next = True
    if text:
        run = p.add_run(text)
        set_run_font(run, face, size, bold=bold, italic=italic, color=color)
    return p

def add_run_to(para, text, face, size, bold=False, italic=False, color=None):
    run = para.add_run(text)
    set_run_font(run, face, size, bold=bold, italic=italic, color=color)
    return run

def add_rich_runs(para, text, face, size, base_bold=False, italic=False, color=None):
    """将 text 按 **bold** 拆段，交替渲染为粗体/正常 runs。"""
    if not text:
        return
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            add_run_to(para, part[2:-2], face, size, bold=True, italic=italic, color=color)
        else:
            add_run_to(para, part, face, size, bold=base_bold, italic=italic, color=color)

def set_cell_font(cell, text, face, size, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                  bg_color=None, text_color=None):
    # 保留 **bold**，先剥离引用标注
    text = clean_text_of_citations(text)
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)
    # 拆分 **...** 内联粗体
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if not part:
            continue
        is_bold_part = part.startswith('**') and part.endswith('**')
        show = part[2:-2] if is_bold_part else part
        run = p.add_run(show)
        run.bold = bold or is_bold_part
        run.font.name = face
        run.font.size = size
        if text_color:
            run.font.color.rgb = RGBColor(*text_color)
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), face)
        rPr.insert(0, rFonts)
    if bg_color:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg_color)
        tcPr.append(shd)

def add_table_to_doc(doc, headers, rows):
    n_cols = len(headers)
    avail_w = PAGESIZE[0] - 2*MARGIN
    col_w = avail_w / n_cols
    tbl = doc.add_table(rows=1+len(rows), cols=n_cols)
    tbl.style = 'Table Grid'
    # 表头
    for ci, h in enumerate(headers):
        set_cell_font(tbl.rows[0].cells[ci], h, HEAD_FONT, Pt(10.5),
                      bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                      bg_color='2F5496', text_color=(255,255,255))
    # 数据行
    for ri, row in enumerate(rows):
        bg = 'D6E4F0' if ri % 2 == 0 else 'FFFFFF'
        # 行数与 header 列数对齐（过长截断，过短补空）
        cells = list(row) + [""] * max(0, n_cols - len(row))
        cells = cells[:n_cols]
        for ci, val in enumerate(cells):
            align = WD_ALIGN_PARAGRAPH.CENTER if ci == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_font(tbl.rows[ri+1].cells[ci], str(val), BODY_FONT, Pt(9.5),
                          align=align, bg_color=bg)
    doc.add_paragraph()  # 空段后距
    return tbl

PAGESIZE = ( Cm(21), Cm(29.7) )  # A4
MARGIN = Cm(3.18)

def set_page_margins(doc):
    for s in doc.sections:
        s.top_margin    = MARGIN
        s.bottom_margin = MARGIN
        s.left_margin  = MARGIN
        s.right_margin = MARGIN

# ── 引用规范化 ────────────────────────────────────────
CITATION_MAP = {}   # citation_key → ref_number
CITATION_COUNTER = [0]

def normalize_citation(cite_str):
    """将 [B级, 来源, 年份] 转为 [1] 格式并记录映射"""
    m = re.search(r'\[(.+?)\]', cite_str)
    if not m:
        return cite_str
    inner = m.group(1)
    # 提取关键来源名（去掉级别和日期）
    parts = [p.strip() for p in inner.split(',')]
    if len(parts) >= 2:
        key = f"{parts[0]}|{parts[1]}"
    else:
        key = inner
    if key not in CITATION_MAP:
        CITATION_COUNTER[0] += 1
        CITATION_MAP[key] = CITATION_COUNTER[0]
    num = CITATION_MAP[key]
    cite_str = cite_str.replace(m.group(0), f"[{num}]")
    return cite_str

def clean_text_of_citations(text):
    """从正文中移除引用标注，保留非引用类括号内容"""
    # 只清除含"级"或媒体/年份格式的引用 [A/B/C级, 来源, 年份]
    text = re.sub(r'\[(?:A|B|C|D)级\s*,\s*[^]]{2,60}\]', '', text)
    # 清除纯年份/数字引用 [YYYY] 或 [数字]
    text = re.sub(r'\[(?:\d{4}年?|\d+)\]', '', text)
    return text.strip()

def strip_markdown(text):
    """清除 markdown 格式符号但保留纯文本"""
    # **bold** → bold
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # *italic* → italic
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    # ### 残留 → 删除
    text = re.sub(r'^#{1,6}\s*', '', text)
    return text

# ── 标题层级体系 ────────────────────────────────────
def h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(24)
    pf.space_after  = Pt(12)
    pf.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, HEAD_FONT, HEAD1_SIZE, bold=True)
    # 下划线装饰
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(12)
    run2 = p2.add_run()
    run2.font.name = HEAD_FONT
    pPr = run2._element.get_or_add_rPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2F5496')
    pBdr.append(bottom)
    pPr.append(pBdr)

def h2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after  = Pt(8)
    pf.keep_with_next = True
    run = p.add_run("■ " + text)
    set_run_font(run, HEAD_FONT, HEAD2_SIZE, bold=True, color=(0,0,128))

def h3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after  = Pt(6)
    pf.keep_with_next = True
    run = p.add_run("● " + text)
    set_run_font(run, HEAD_FONT, HEAD3_SIZE, bold=True, color=(80,80,80))

def body(doc, text):
    """正文段落：首行缩进，1.2倍行距，仿宋四号"""
    text = clean_text_of_citations(text)           # 剥离引用但保留 **bold**
    if not text:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = LINE_SPACING
    pf.first_line_indent = INDENT_EM
    add_rich_runs(p, text, BODY_FONT, BODY_SIZE)

def quote_block(doc, text):
    text = text.strip()
    if not text:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after  = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = LINE_SPACING
    pf.left_indent = Cm(1.0)
    run = p.add_run(text)
    run.font.name  = BODY_FONT
    run.font.size  = Pt(12)
    run.italic     = True
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), BODY_FONT)
    rPr.insert(0, rFonts)

# ── Markdown 解析 ────────────────────────────────────
def process_md(doc, lines, person_name="政治人物"):
    tbl_hdr, tbl_rows, in_tbl = [], [], False
    i = 0

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        i += 1

        # ── 跳过全空行（表格外） ──
        if not line.strip():
            if not in_tbl:
                pass  # 空行不占位
            continue

        # ── 版本头注释块 → 标题页 ──
        if line.startswith('**报告版本**') and 'OSINT' in line:
            meta = []
            while i < len(lines) and lines[i].strip() and not lines[i].startswith('#'):
                meta.append(lines[i].rstrip())
                i += 1
            # 标题
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_title.paragraph_format.space_before = Pt(72)
            p_title.paragraph_format.space_after  = Pt(24)
            r = p_title.add_run(f"{person_name} OSINT 政治人物全景分析报告")
            set_run_font(r, HEAD_FONT, Pt(22), bold=True)

            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.paragraph_format.space_after = Pt(48)
            r2 = p_sub.add_run("（OSINT V7.2 全景分析体系）")
            set_run_font(r2, HEAD_FONT, Pt(14))

            doc.add_paragraph()
            for ml in meta:
                m = re.match(r'\*\*(.+?)\*\*[:：]\s*(.+)', ml.strip())
                if m:
                    k, v = m.groups()
                    k = strip_markdown(k.strip())   # 清除 **bold** 残留
                    v = strip_markdown(v.strip())
                    pm = doc.add_paragraph()
                    pm.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    pm.paragraph_format.space_before = Pt(2)
                    pm.paragraph_format.space_after  = Pt(2)
                    add_run_to(pm, f"{k}：", HEAD_FONT, Pt(12), bold=True)
                    add_run_to(pm, v, BODY_FONT, Pt(12))
            doc.add_page_break()
            continue

        # ── 跳过 markdown 标题行本身 ──
        if re.match(r'^#{1,3}\s+[^\s]', line) and not line.startswith('## '):
            continue

        # ── 一级标题 ## ──
        if line.startswith('## '):
            if in_tbl:
                add_table_to_doc(doc, tbl_hdr, tbl_rows)
                tbl_hdr, tbl_rows, in_tbl = [], [], False
            text = line[3:].strip()
            h1(doc, text)
            continue

        # ── 二级标题 ### ──
        if line.startswith('### '):
            if in_tbl:
                add_table_to_doc(doc, tbl_hdr, tbl_rows)
                tbl_hdr, tbl_rows, in_tbl = [], [], False
            text = line[4:].strip()
            # 三级模块标题用 ●
            if text.startswith('§') or re.match(r'模块\s*\d+', text) or '矩阵' in text or '量化' in text or '核查' in text:
                h3(doc, text)
            else:
                h2(doc, text)
            continue

        # ── 引用 ──
        if line.startswith('>'):
            quote_block(doc, line[1:].strip())
            continue

        # ── 表格 ──
        if line.startswith('|'):
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                continue
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if not in_tbl:
                in_tbl = True
                tbl_hdr = cells
            else:
                tbl_rows.append(cells)
            continue
        else:
            if in_tbl:
                add_table_to_doc(doc, tbl_hdr, tbl_rows)
                tbl_hdr, tbl_rows, in_tbl = [], [], False

        # ── 正文 ──
        body(doc, line)

    if in_tbl:
        add_table_to_doc(doc, tbl_hdr, tbl_rows)

# ── 主程序 ────────────────────────────────────────────
def main():
    print("📖 读取 markdown...")
    with open(SRC, encoding='utf-8') as f:
        raw = f.read()
    lines = raw.splitlines()

    # 从第一行提取人名（格式：# 蒋万安 OSINT...）
    m = re.match(r'^#\s+([^\s]+)\s+OSINT', lines[0] if lines else '')
    person_name = m.group(1) if m else "政治人物"
    print(f"   检测到人物：{person_name}")

    print("📝 创建 Word 文档...")
    doc = Document()
    set_page_margins(doc)

    print("⚙️  处理内容...")
    process_md(doc, lines, person_name)

    if OUT_DOCX:
        out = OUT_DOCX
    else:
        base = os.path.splitext(os.path.basename(SRC))[0]
        out = os.path.join(OUT_DIR, base + ".docx")
    doc.save(out)
    sz = os.path.getsize(out)
    print(f"✅ DOCX 已保存：{out}")
    print(f"   大小：{sz:,} bytes ({sz/1024:.1f} KB)")

if __name__ == '__main__':
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", help="source markdown path")
    ap.add_argument("--out", help="output docx path")
    args = ap.parse_args()
    if args.src:
        SRC = args.src
        OUT_DIR = os.path.dirname(args.src)
    if args.out:
        OUT_DOCX = args.out
    main()
