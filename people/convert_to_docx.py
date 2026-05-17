"""
蒋万安 OSINT 政治人物分析报告 v1.8 → Word (.docx) 转换器
中文期刊论文格式规范
"""
import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 路径配置 ──────────────────────────────────────────
SRC = r"D:\aicoding\news-monitor\people\output\蒋万安\蒋万安_政治人物分析报告_v1.8.md"
OUT_DIR = r"D:\aicoding\news-monitor\people\output\蒋万安"

# ── 中文学术期刊格式常量 ────────────────────────────────
FONT_BODY   = "宋体"    # 正文
FONT_HEAD   = "黑体"    # 一级标题
FONT_SUB    = "楷体"    # 二级标题/副标题
FONT_CODE   = "Courier New"  # 代码/表格
MARGIN_TOP    = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)
MARGIN_LEFT   = Cm(3.0)
MARGIN_RIGHT  = Cm(3.0)
BODY_FONT_SIZE = Pt(12)    # 正文12pt
HEAD_FONT_SIZE = Pt(15)    # 一级标题15pt
SUB_FONT_SIZE  = Pt(13)    # 二级标题13pt

# ── helpers ────────────────────────────────────────────
def set_margins(doc):
    for section in doc.sections:
        section.top_margin    = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin  = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT

def add_run(para, text, bold=False, italic=False, font_name=FONT_BODY,
            font_size=Pt(12), color=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font_name
    run.font.size = font_size
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)
    return run

def new_paragraph(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  bold=False, font_name=FONT_BODY, font_size=BODY_FONT_SIZE,
                  space_before=Pt(0), space_after=Pt(6), line_spacing=1.5,
                  color=None, keep_with_next=False):
    para = doc.add_paragraph()
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = space_before
    pf.space_after  = space_after
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing      = line_spacing
    if keep_with_next:
        pf.keep_with_next = True
    if text:
        add_run(para, text, bold=bold, font_name=font_name,
                font_size=font_size, color=color)
    return para

def add_heading(doc, text, level=1):
    """中文期刊风格标题"""
    if level == 1:
        # 一级：■ 居中，黑体15pt
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = para.paragraph_format
        pf.space_before = Pt(18)
        pf.space_after  = Pt(6)
        pf.keep_with_next = True
        add_run(para, text, bold=True, font_name=FONT_HEAD,
                font_size=Pt(16), color=(0, 0, 0))
        return para
    elif level == 2:
        # 二级：■ 左对齐，黑体14pt
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after  = Pt(4)
        pf.keep_with_next = True
        add_run(para, text, bold=True, font_name=FONT_HEAD,
                font_size=Pt(14), color=(0, 0, 0))
        return para
    elif level == 3:
        # 三级：■ 左对齐，楷体12pt
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = para.paragraph_format
        pf.space_before = Pt(8)
        pf.space_after  = Pt(3)
        pf.keep_with_next = True
        add_run(para, text, bold=True, font_name=FONT_SUB,
                font_size=Pt(12), color=(0, 0, 0))
        return para

def add_table(doc, header, rows, col_widths=None):
    """添加表格，处理中文字体"""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = 'Table Grid'
    # 表头
    hrow = table.rows[0]
    for i, h in enumerate(header):
        cell = hrow.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), FONT_BODY)
        rPr.insert(0, rFonts)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2F5496')
        tcPr.append(shd)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 数据行
    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        fill_color = 'D6E4F0' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = FONT_BODY
            run.font.size = Pt(9)
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), FONT_BODY)
            rPr.insert(0, rFonts)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill_color)
            tcPr.append(shd)
    doc.add_paragraph()  # 空行
    return table

# ── markdown parser ────────────────────────────────────
def parse_markdown(src):
    with open(src, encoding='utf-8') as f:
        lines = f.readlines()
    return lines

def process_lines(doc, lines):
    i = 0
    in_table = False
    table_header = []
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip('\n\r')
        i += 1

        # 空行
        if not line.strip():
            if in_table:
                # flush table
                if table_header and table_rows:
                    add_table(doc, table_header, table_rows)
                table_header = []
                table_rows = []
                in_table = False
            continue

        # 跳过版本头注释块
        if line.startswith('**报告版本**') and 'OSINT' in line:
            # 读取报告基本信息块
            info_lines = []
            while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and '##' not in lines[i][:5]:
                info_lines.append(lines[i].rstrip())
                i += 1
            # 添加标题页
            add_title_page(doc, info_lines)
            continue

        # 跳过独立行 markdown 标题（已在标题页处理）
        if line.startswith('# 蒋万安 OSINT'):
            continue

        # 一级标题 ■ ■ ■
        if line.startswith('## '):
            text = line[3:].strip()
            add_heading(doc, text, level=1)
            continue

        # 二级标题 ■ ■
        if line.startswith('### '):
            text = line[4:].strip()
            # 特殊处理表格标题标记
            if text.startswith('§') or '模块' in text or '矩阵' in text:
                add_heading(doc, text, level=3)
            else:
                add_heading(doc, text, level=2)
            continue

        # 引用块（> 开头）
        if line.startswith('>'):
            text = line[1:].strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = para.paragraph_format
            pf.left_indent = Cm(0.8)
            pf.space_before = Pt(3)
            pf.space_after  = Pt(3)
            run = para.add_run(text)
            run.italic = True
            run.font.name = FONT_SUB
            run.font.size = Pt(11)
            r = run._element
            rPr = r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), FONT_SUB)
            rPr.insert(0, rFonts)
            continue

        # 表格行
        if line.startswith('|'):
            # 检查是否是表头行
            if '---' in line or all(c in '-|:' for c in line):
                i += 1  # 跳过分隔行
                continue
            cells = [c.strip() for c in line.strip('|').split('|')]
            if not in_table:
                in_table = True
                table_header = cells
            else:
                table_rows.append(cells)
            continue
        else:
            if in_table:
                if table_header and table_rows:
                    add_table(doc, table_header, table_rows)
                table_header = []
                table_rows = []
                in_table = False

        # 普通正文
        # 检测加粗关键词（来源标注）
        line_clean = line.strip()
        if not line_clean:
            continue

        # 处理行内加粗
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', line_clean)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing      = 1.5

        has_content = False
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                text = part[2:-2]
                add_run(para, text, bold=True, font_name=FONT_BODY,
                        font_size=BODY_FONT_SIZE)
                has_content = True
            elif part.startswith('*') and part.endswith('*'):
                text = part[1:-1]
                add_run(para, text, italic=True, font_name=FONT_SUB,
                        font_size=BODY_FONT_SIZE)
                has_content = True
            elif part:
                add_run(para, part, font_name=FONT_BODY, font_size=BODY_FONT_SIZE)
                has_content = True

        if not has_content:
            # 空段落，但保持格式
            add_run(para, '', font_name=FONT_BODY, font_size=BODY_FONT_SIZE)

    # flush final table
    if in_table and table_header and table_rows:
        add_table(doc, table_header, table_rows)

def add_title_page(doc, info_lines):
    """生成中文期刊风格标题页"""
    # 报告标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = title_para.paragraph_format
    pf.space_before = Pt(72)  # 大留白
    pf.space_after  = Pt(24)
    title_run = title_para.add_run('蒋万安 OSINT 政治人物全景分析报告')
    title_run.bold = True
    title_run.font.name = FONT_HEAD
    title_run.font.size = Pt(22)
    r = title_run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), FONT_HEAD)
    rPr.insert(0, rFonts)

    # 副标题
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf2 = sub_para.paragraph_format
    pf2.space_before = Pt(0)
    pf2.space_after  = Pt(48)
    # 提取关键元信息
    meta_info = {}
    for line in info_lines:
        m = re.match(r'\*\*(.+?)\*\*[:：]\s*(.+)', line.strip())
        if m:
            key, val = m.groups()
            meta_info[key.strip()] = val.strip()

    # 元信息块
    meta_keys = ['分析对象', '当前职务', '编制日期', '保密级别', '行为框架', '数据采集', '版本说明']
    for key, val in meta_info.items():
        if key in meta_keys:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf3 = meta_para.paragraph_format
            pf3.space_before = Pt(3)
            pf3.space_after  = Pt(3)
            pf3.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf3.line_spacing = 1.3
            add_run(meta_para, f'{key}：', bold=True, font_name=FONT_SUB,
                    font_size=Pt(11))
            add_run(meta_para, val, font_name=FONT_BODY, font_size=Pt(11))

    # 副标题（框架版本）
    framework = meta_info.get('框架版本', 'OSINT V7.2')
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf2 = sub_para.paragraph_format
    pf2.space_before = Pt(0)
    pf2.space_after  = Pt(48)
    sub_run = sub_para.add_run(f'（{framework}）')
    sub_run.font.name = FONT_SUB
    sub_run.font.size = Pt(14)
    r2 = sub_run._element
    rPr2 = r2.get_or_add_rPr()
    rFonts2 = OxmlElement('w:rFonts')
    rFonts2.set(qn('w:eastAsia'), FONT_SUB)
    rPr2.insert(0, rFonts2)

    # 分页
    doc.add_page_break()

# ── 主程序 ────────────────────────────────────────────
def main():
    print("📖 读取 markdown 文件...")
    lines = parse_markdown(SRC)

    print("📝 创建 Word 文档...")
    doc = Document()
    set_margins(doc)

    print("⚙️  处理内容...")
    process_lines(doc, lines)

    out_docx = os.path.join(OUT_DIR, '蒋万安_政治人物分析报告_v1.8.docx')
    doc.save(out_docx)
    size = os.path.getsize(out_docx)
    print(f"✅ Word 文档已保存：{out_docx}")
    print(f"   文件大小：{size:,} bytes ({size/1024:.1f} KB)")
    return out_docx

if __name__ == '__main__':
    main()
