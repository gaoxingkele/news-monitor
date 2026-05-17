from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle


TITLE = "郑丽文来访大陆话题过去48小时社交媒体热帖归类"
SUBTITLE = "时间窗口：2026-04-07 至 2026-04-09"

OVERVIEW = [
    "过去48小时内，关于“郑丽文来访大陆/郑习会”的社交媒体热帖，呈现出明显的平台分层。国际传播层面主要由 X/Twitter 带动声量，且传播最强的帖子并不全是负面批评，很多是把此事包装为“和平之旅”“历史性访问”的中性或正向信息；而中文论坛层面则更容易迅速转化为‘亲中/卖台/统战/军购/选战’等高度冲突化叙事。",
    "从可获取性和证据强度看，本轮最有价值的平台依次是：X/Twitter、Reddit、Dcard、PTT、Mobile01。LIHKG、Facebook 公共页、巴哈姆特在这个时间窗内缺少足够强的可核验高热样本，不宜作为主证据源。",
    "总体情绪判断不是单边倒，而是‘分裂型’：国际平台上的高传播帖更容易出现“和平沟通/历史性动作”的包装；台湾年轻平台和政治论坛则更多把此事视为统战样板、路线站队和岛内政争素材。"
]

PLATFORM_SECTIONS = [
    {
        "title": "一、X/Twitter：国际传播主战场",
        "paragraphs": [
            "X/Twitter 是过去48小时里最强的传播平台。它同时容纳了主流媒体账号、地缘政治账号、亲中评论账号和反中评论账号，因此同一事件在这里表现为高度分裂，但也最容易看出传播量级与议题扩散方向。",
            "本轮高传播帖里，传播量最高的是将郑丽文来访大陆定义为“和平之旅”“历史性会面”的帖子，而不是单纯批评帖。与此同时，反对声量主要集中在若干台湾/海外评论账号，重点攻击点是‘亲中’‘卖台’‘为北京服务’。"
        ],
        "items": [
            ("ABC", "中性报道，强调这是应习近平邀请的“和平之旅”，传播量高。约 331 likes / 86 reposts / 73 replies / 101k views。", "https://x.com/ABC/status/2041473564742684966"),
            ("Carl Zha", "传播量最高的一组正向叙事账号，突出“十年来首次KMT主席访中”“历史性时刻”。4月7日主帖约 8k likes / 1.3k reposts / 408k views。", "https://x.com/CarlZha/status/2041481953766203445"),
            ("Carl Zha（二次传播）", "4月8日继续放大“南京中山陵+赴京会习”的仪式化意义。约 3.4k likes / 465 reposts / 84k views。", "https://x.com/CarlZha/status/2041788238319513732"),
            ("Ian Miles Cheong", "将此事直接解释为台湾向中国靠近、统一路线升温。约 888 likes / 277 reposts / 120k views。", "https://x.com/ianmiles/status/2041904586387312717"),
            ("denisewu", "明确负面，把郑丽文比附为亲北京、牺牲台湾利益的政治人物。约 104 likes / 14 reposts / 7k views。", "https://x.com/denisewu/status/2041369740359377294"),
            ("clashreport", "中性放大“十年来首次KMT领袖访中”的新闻性和地缘政治意味。约 404 likes / 52 reposts / 36k views。", "https://x.com/clashreport/status/2041402814908268663"),
        ],
    },
    {
        "title": "二、Reddit：英文圈的路线与安全争论",
        "paragraphs": [
            "Reddit 的热度不如 X 高，但讨论质量更集中于‘台美防务合作是否会受影响’‘KMT 是否进一步靠近北京’以及‘维持现状与统一之间是否还有灰区’。",
            "Reddit 的评论区整体偏负面或偏怀疑，尤其在 r/worldnews 这类英文公共区里，用户更容易从香港经验、统一风险和 KMT 历史定位出发进行批评。"
        ],
        "items": [
            ("r/worldnews 热帖 1", "焦点是‘访问大陆是否会削弱台湾安全与台美协作’，评论区高频词包括 Hong Kong、KMT、sellout、status quo。", "https://www.reddit.com/r/worldnews/comments/1sdnt9j/taiwan_opposition_leader_to_visit_china_as/"),
            ("r/worldnews 热帖 2", "过去48小时内的持续热帖，评论区更情绪化，典型表达包括“they can keep her”“KMT stooges”。", "https://www.reddit.com/r/worldnews/comments/1sexjak/taiwans_opposition_leader_arrives_in_china_for_a/"),
            ("r/worldnews 热帖 3", "较多从历史讽刺角度批评KMT路线：从反共到更公开靠近北京。", "https://www.reddit.com/r/worldnews/comments/1serrde/cheng_liwun_taiwan_opposition_kuomintang_leader/"),
            ("r/China 背景热帖", "更系统地讨论‘阻军购预算’‘对北京示好’‘绕开赖政府’等问题。", "https://www.reddit.com/r/China/comments/1s80s72/taiwan_opposition_leader_accepts_xis_invite_to/"),
        ],
    },
    {
        "title": "三、台湾论坛：情绪比事实跑得更快",
        "paragraphs": [
            "台湾论坛层面，过去48小时最值得看的不是单条新闻，而是平台间的差异。PTT 最情绪化、最选战化；Dcard 更接近年轻用户对身份认同和路线风险的即时反应；Mobile01 则更容易出现反绿、反‘安全叙事操作’的逆向声音。",
            "也就是说，同一事件在台湾论坛上不是一个公共共识议题，而是被拆成不同社群的政治身份表达。"
        ],
        "items": [
            ("Dcard：郑丽文参访大陆", "偏负面，代表年轻用户对时机与路线的警惕，常见说法是‘现在访中只会让人更退缩’。", "https://www.dcard.tw/f/talk/p/261254296"),
            ("Dcard：真的假的，郑丽文直接公开高喊中华民国", "相对少见的正向/辩护型帖子，试图证明其并非完全向北京让步。", "https://www.dcard.tw/f/trending/p/261259450"),
            ("Dcard：郑丽文打脸赖政府", "代表反绿、反赖政府操作的逆向声量。", "https://www.dcard.tw/f/trending/p/261251055"),
            ("PTT HatePolitics", "不是新帖但仍是讨论入口，关键词集中在‘卖台/抹红/加分减分/2028路线战’。", "https://www.ptt.cc/bbs/HatePolitics/M.1774838712.A.083.html"),
            ("Mobile01：国共交流北京主导", "偏警戒，强调此行背后是北京设定议程与统战推进。", "https://www.mobile01.com/topicdetail.php?f=638&t=7248718"),
            ("Mobile01：中网灌爆“臭乞丐”", "补到大陆网民侧的嘲讽与敌意反应，但其本身仍是媒体转述而非原生大陆平台抓取。", "https://www.mobile01.com/topicdetail.php?f=638&p=1&t=7244919"),
            ("Mobile01：简舒培批“投名状”", "代表绿营支持者的攻击框架在论坛继续扩散。", "https://www.mobile01.com/topicdetail.php?f=638&t=7245404"),
        ],
    },
    {
        "title": "四、证据较弱的平台",
        "paragraphs": [
            "LIHKG：这48小时内没有找到足够强的新热串。可检到的相关内容更多是郑丽文长期身份认同争议的背景串，而不是本轮事件的直接高热讨论。",
            "Facebook 公共页/公共贴文：搜索索引可见度较差，缺少足够稳定且可公开验证互动量的高热贴，不适合作为主证据源。",
            "巴哈姆特：没有检到与本话题直接相关且讨论量足够的公开高热串。"
        ],
        "items": [
            ("LIHKG 背景串", "提供郑丽文长期标签化认知的背景情绪参考，但不是过去48小时主战场。", "https://lihkg.com/thread/3979309/page/1"),
        ],
    },
]

NARRATIVES = [
    "和平之旅/对话降温：在国际传播层面声量很强，尤其是中性媒体和亲中账号，倾向把此行包装为降低风险、恢复沟通的动作。",
    "亲中/卖台/统战样板：在中文社媒、Reddit评论区和台湾论坛里最强，很多用户不接受‘和平之旅’定义，而是把此行直接理解成向北京靠拢。",
    "影响军购与台美关系：英文圈和较理性的政治讨论常把访问与台湾防务预算、对美安全合作联系起来。",
    "历史性象征 / KMT路线再定义：不少国际与评论账号强调‘十年来首次在任KMT主席访中’，把这件事理解为路线信号而不是单次访问。",
    "岛内选战化 / 蓝绿路线战：在 PTT、Dcard、Mobile01 上非常明显，很多讨论最终会落到‘谁加分、谁卖台、谁在操作’。",
]

CONCLUSION = [
    "过去48小时，这个话题在社交媒体上的高热并不是简单的单向支持或反对，而是高度分裂。",
    "从传播量看，最强的单帖往往来自中性媒体或正向包装“和平之旅”的账号；从情绪强度看，最尖锐的批评则主要集中在中文论坛、评论账号和 Reddit 评论区。",
    "如果后续要持续监测这一话题，最值得盯住的平台顺序是：X/Twitter、Reddit、Dcard、PTT、Mobile01。它们已经足以覆盖国际传播、英文圈态度、台湾年轻用户、政治论坛和逆向声量。"
]


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_docx(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title = doc.add_heading(TITLE, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(SUBTITLE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    doc.add_heading("总体概览", level=1)
    for text in OVERVIEW:
        doc.add_paragraph(text)

    for section in PLATFORM_SECTIONS:
        doc.add_heading(section["title"], level=1)
        for text in section["paragraphs"]:
            doc.add_paragraph(text)
        for name, summary, url in section["items"]:
            para = doc.add_paragraph()
            para.add_run(f"{name}：").bold = True
            para.add_run(summary + " 链接：")
            add_hyperlink(para, url, url)

    doc.add_heading("五、主要叙事归类", level=1)
    for text in NARRATIVES:
        doc.add_paragraph(text)

    doc.add_heading("六、结论", level=1)
    for text in CONCLUSION:
        doc.add_paragraph(text)

    doc.save(output_path)


def build_pdf(output_path: Path) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleCN", parent=styles["Title"], fontName="STSong-Light", fontSize=18, leading=24, spaceAfter=12)
    body_style = ParagraphStyle("BodyCN", parent=styles["BodyText"], fontName="STSong-Light", fontSize=11, leading=16, spaceAfter=6)
    heading_style = ParagraphStyle("HeadingCN", parent=styles["Heading1"], fontName="STSong-Light", fontSize=14, leading=18, spaceBefore=8, spaceAfter=8)

    story = [
        Paragraph(TITLE, title_style),
        Paragraph(SUBTITLE, body_style),
        Paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}", body_style),
        Spacer(1, 0.3 * cm),
        Paragraph("总体概览", heading_style),
    ]
    for text in OVERVIEW:
        story.append(Paragraph(text, body_style))

    for section in PLATFORM_SECTIONS:
        story.append(Paragraph(section["title"], heading_style))
        for text in section["paragraphs"]:
            story.append(Paragraph(text, body_style))
        for name, summary, url in section["items"]:
            story.append(Paragraph(f"<b>{name}：</b>{summary}<br/><font color='blue'>{url}</font>", body_style))

    story.append(Paragraph("五、主要叙事归类", heading_style))
    for text in NARRATIVES:
        story.append(Paragraph(text, body_style))
    story.append(Paragraph("六、结论", heading_style))
    for text in CONCLUSION:
        story.append(Paragraph(text, body_style))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=2.2 * cm,
        rightMargin=2.2 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
    )
    doc.build(story)


def main() -> None:
    out_dir = Path(r"D:\BaiduSyncdisk\2026-04\台情")
    out_dir.mkdir(parents=True, exist_ok=True)
    base = "郑丽文来访大陆过去48小时社交热帖归类_20260409"
    docx_path = out_dir / f"{base}.docx"
    pdf_path = out_dir / f"{base}.pdf"
    build_docx(docx_path)
    build_pdf(pdf_path)
    print(docx_path)
    print(pdf_path)


if __name__ == "__main__":
    main()
