"""把 2026-04-24 的所有最终产出打包到 topics/0424/ 下，含 md/docx/pdf/xlsx。

目录结构：
  topics/0424/
    ├─ reports/                 核心报告 md+docx+pdf（分析/项目状态/可疑详单/精修摘要）
    ├─ xlsx/                    合并总表 + 13 份批次 refined + targets JSON
    ├─ charts/                  6 张对比图 PNG
    ├─ per_country/md/          64 份国家独立 MD
    ├─ per_country/docx/        64 份国家独立 DOCX
    └─ per_country/pdf/         64 份国家独立 PDF
"""
from __future__ import annotations
import io, os, re, shutil, sys, warnings
from datetime import datetime
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
ROOT = Path(r"D:\aicoding\news-monitor")
TOPICS = ROOT / "topics"
DATE = "20260424"
OUT = TOPICS / "0424"

FONT_PATH = "C:/Windows/Fonts/msyh.ttc"   # 微软雅黑

from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT


# ─── MD → PDF ─────────────────────────────────────────────────────────────
def md_to_pdf(md_path: Path, pdf_path: Path, charts_root: Path = None) -> None:
    content = md_path.read_text(encoding="utf-8")
    lines = content.split("\n")

    pdf = FPDF()
    pdf.set_margins(12, 12, 12)
    pdf.set_auto_page_break(auto=True, margin=12)
    with warnings.catch_warnings():
        warnings.simplefilter("ignore", DeprecationWarning)
        try: pdf.add_font("CJK", "", FONT_PATH, uni=True)
        except TypeError: pdf.add_font("CJK", "", FONT_PATH)
    pdf.add_page()

    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False; return
        n_cols = max(len(r) for r in table_rows)
        usable = pdf.w - pdf.l_margin - pdf.r_margin
        col_w = usable / n_cols if n_cols else usable
        if table_rows:
            pdf.set_font("CJK", size=7)
            pdf.set_fill_color(230, 230, 230)
            for cell in table_rows[0]:
                pdf.cell(col_w, 5, cell.strip()[:30], border=1, fill=True, align="C")
            pdf.ln()
        pdf.set_fill_color(255, 255, 255)
        for row in table_rows[1:]:
            if all(re.match(r"^[:\-\s]*$", c) for c in row): continue
            for i, cell in enumerate(row):
                align = "L" if i == 0 else "C"
                pdf.cell(col_w, 5, cell.strip()[:30], border=1, align=align)
            pdf.ln()
        pdf.ln(2)
        in_table = False; table_rows = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if all(re.match(r"^[:\-]+$", c) for c in cells):
                in_table = True; continue
            table_rows.append(cells); in_table = True; continue
        elif in_table:
            flush_table()

        # ![alt](path) 图片
        img_match = re.match(r'^!\[[^\]]*\]\(([^)]+)\)$', stripped)
        if img_match and charts_root:
            rel = img_match.group(1)
            img_path = charts_root / Path(rel).name
            if img_path.exists():
                try:
                    pdf.ln(2)
                    pdf.image(str(img_path), w=pdf.w - pdf.l_margin - pdf.r_margin - 10)
                    pdf.ln(3)
                except Exception as e:
                    pdf.multi_cell(0, 5, f"[图片加载失败：{img_path.name}] {e}",
                                   new_x="LMARGIN", new_y="NEXT")
                continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            pdf.set_font("CJK", size=15)
            pdf.multi_cell(0, 9, stripped[2:], align="C", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)
        elif stripped.startswith("## "):
            pdf.ln(2); pdf.set_font("CJK", size=12); pdf.set_fill_color(240, 240, 240)
            pdf.multi_cell(0, 7, stripped[3:], align="L", new_x="LMARGIN", new_y="NEXT", fill=True)
            pdf.ln(2)
        elif stripped.startswith("### "):
            pdf.ln(1); pdf.set_font("CJK", size=10)
            pdf.multi_cell(0, 6, stripped[4:], align="L", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
        elif stripped.startswith("#### "):
            pdf.set_font("CJK", size=9)
            pdf.multi_cell(0, 5.5, stripped[5:], align="L", new_x="LMARGIN", new_y="NEXT")
        elif stripped.startswith("---"):
            y = pdf.get_y(); pdf.set_draw_color(180, 180, 180)
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y); pdf.ln(2)
        elif stripped == "":
            pdf.ln(1)
        elif stripped.startswith("- "):
            pdf.set_font("CJK", size=8); pdf.set_x(pdf.l_margin + 3)
            pdf.multi_cell(0, 4.5, "• " + stripped[2:], new_x="LMARGIN", new_y="NEXT")
        elif stripped.startswith("  - "):
            pdf.set_font("CJK", size=7); pdf.set_x(pdf.l_margin + 6)
            pdf.multi_cell(0, 4.2, "◦ " + stripped[4:], new_x="LMARGIN", new_y="NEXT")
        elif stripped.startswith("  http") or stripped.startswith("    http"):
            pdf.set_font("CJK", size=7); pdf.set_text_color(30, 100, 200)
            pdf.set_x(pdf.l_margin + 6)
            pdf.multi_cell(0, 4, stripped, new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)
        elif stripped.startswith("> "):
            pdf.set_font("CJK", size=8); pdf.set_text_color(80, 80, 80)
            pdf.set_x(pdf.l_margin + 5)
            pdf.multi_cell(0, 4.5, stripped[2:], new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)
        else:
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            pdf.set_font("CJK", size=9)
            pdf.multi_cell(0, 5, text, new_x="LMARGIN", new_y="NEXT")

    if in_table: flush_table()
    pdf.output(str(pdf_path))


# ─── MD → DOCX ────────────────────────────────────────────────────────────
def set_run_font(run, face="微软雅黑", size=Pt(11), bold=False, color=None):
    run.bold = bold
    run.font.name = face
    run.font.size = size
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), face)
    rPr.insert(0, rFonts)


def add_para(doc, text, *, size=Pt(11), bold=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, color=None,
             space_before=Pt(0), space_after=Pt(4), indent_cm=0):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    if indent_cm:
        pf.first_line_indent = Cm(indent_cm)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_table_from_rows(doc, rows, charts_root=None):
    if not rows: return
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = tbl.cell(i, j)
            text = row[j] if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            set_run_font(run, size=Pt(8), bold=(i == 0))


def md_to_docx(md_path: Path, docx_path: Path, charts_root: Path = None) -> None:
    content = md_path.read_text(encoding="utf-8")
    lines = content.split("\n")

    doc = Document()
    # 默认样式
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    rPr.insert(0, rFonts)
    for sec in doc.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)

    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False; return
        # 过滤分隔行
        filtered = [r for r in table_rows
                    if not all(re.match(r'^[:\-]+$', c.strip()) for c in r)]
        add_table_from_rows(doc, filtered, charts_root=charts_root)
        doc.add_paragraph()
        in_table = False; table_rows = []

    for line in lines:
        stripped = line.strip()

        # 表格
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if all(re.match(r'^[:\-]+$', c) for c in cells):
                in_table = True; continue
            table_rows.append(cells); in_table = True; continue
        elif in_table:
            flush_table()

        # 图片
        img_match = re.match(r'^!\[[^\]]*\]\(([^)]+)\)$', stripped)
        if img_match and charts_root:
            rel = img_match.group(1)
            img_path = charts_root / Path(rel).name
            if img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(6.2))
                except Exception:
                    add_para(doc, f"[图片：{img_path.name}]", size=Pt(9), color=(128,128,128))
                continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            add_para(doc, stripped[2:], size=Pt(20), bold=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     space_before=Pt(0), space_after=Pt(12))
        elif stripped.startswith("## "):
            add_para(doc, stripped[3:], size=Pt(15), bold=True,
                     space_before=Pt(12), space_after=Pt(6), color=(46, 92, 138))
        elif stripped.startswith("### "):
            add_para(doc, stripped[4:], size=Pt(13), bold=True,
                     space_before=Pt(8), space_after=Pt(4))
        elif stripped.startswith("#### "):
            add_para(doc, stripped[5:], size=Pt(11), bold=True,
                     space_before=Pt(4), space_after=Pt(2))
        elif stripped.startswith("---"):
            add_para(doc, "─" * 30, size=Pt(9), color=(180,180,180),
                     align=WD_ALIGN_PARAGRAPH.CENTER)
        elif stripped == "":
            doc.add_paragraph()
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(stripped[2:])
            set_run_font(run, size=Pt(10))
        elif stripped.startswith("  - "):
            p = doc.add_paragraph(style='List Bullet 2')
            run = p.add_run(stripped[4:])
            set_run_font(run, size=Pt(9))
        elif stripped.startswith("> "):
            add_para(doc, stripped[2:], size=Pt(10), color=(80,80,80), indent_cm=0.5)
        elif stripped.startswith("**") and stripped.endswith("**") and stripped.count("**") == 2:
            add_para(doc, stripped.strip("*"), size=Pt(11), bold=True)
        else:
            # 移除 **粗体** 标记保留文本
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            if text:
                add_para(doc, text, size=Pt(10.5))

    if in_table: flush_table()
    doc.save(str(docx_path))


# ─── 打包流水线 ───────────────────────────────────────────────────────────
def main():
    print(f"输出目录: {OUT}")
    OUT.mkdir(parents=True, exist_ok=True)
    (OUT / "reports").mkdir(exist_ok=True)
    (OUT / "xlsx").mkdir(exist_ok=True)
    (OUT / "charts").mkdir(exist_ok=True)
    (OUT / "per_country" / "md").mkdir(parents=True, exist_ok=True)
    (OUT / "per_country" / "docx").mkdir(parents=True, exist_ok=True)
    (OUT / "per_country" / "pdf").mkdir(parents=True, exist_ok=True)

    # 1. 复制 charts
    src_charts = TOPICS / "bri_charts"
    for p in src_charts.glob("*.png"):
        shutil.copy2(p, OUT / "charts" / p.name)
    print(f"✓ charts: {len(list((OUT/'charts').glob('*.png')))} 张")

    # 2. 复制 xlsx（合并表 + 13 份 refined + 原始 v_final 作备份）
    xlsx_files = [
        TOPICS / f"bri_64country_merged_{DATE}.xlsx",
    ]
    xlsx_files += list(TOPICS.glob(f"*_refined447_{DATE}.xlsx"))
    for fp in xlsx_files:
        if fp.exists():
            shutil.copy2(fp, OUT / "xlsx" / fp.name)
    # 目标 JSON + ckpt 也附上
    for p in [ROOT / "output" / f"bri_421_suspicious_targets_{DATE}.json",
              ROOT / "output" / "bri_refine447_ckpt.json"]:
        if p.exists():
            shutil.copy2(p, OUT / "xlsx" / p.name)
    print(f"✓ xlsx: {len(list((OUT/'xlsx').glob('*')))} 个文件")

    # 3. 核心报告（4 份）: md 已在 topics，md/docx/pdf 三格式输出到 reports/
    core_reports = [
        (TOPICS / f"bri_analysis_report_{DATE}.md",        "分析报告_对比分析"),
        (TOPICS / f"bri_64country_project_status_{DATE}.md", "项目状态总览"),
        (TOPICS / f"bri_421_suspicious_detail_{DATE}.md",   "可疑格详细清单"),
        (TOPICS / f"bri_refine447_summary_{DATE}.md",       "精修回写摘要"),
    ]
    for src, label in core_reports:
        if not src.exists():
            print(f"  跳过（文件缺失）: {src.name}")
            continue
        md_dst = OUT / "reports" / src.name
        docx_dst = OUT / "reports" / f"{src.stem}.docx"
        pdf_dst = OUT / "reports" / f"{src.stem}.pdf"
        shutil.copy2(src, md_dst)
        try:
            md_to_docx(src, docx_dst, charts_root=OUT / "charts")
            print(f"  ✓ {label} docx → {docx_dst.name}")
        except Exception as e:
            print(f"  ✗ {label} docx 失败: {e}")
        try:
            md_to_pdf(src, pdf_dst, charts_root=OUT / "charts")
            print(f"  ✓ {label} pdf  → {pdf_dst.name}")
        except Exception as e:
            print(f"  ✗ {label} pdf 失败: {e}")

    # 4. 每国独立文件 64 份（md 复制 + docx + pdf）
    per_country_src = TOPICS / "per_country"
    md_files = sorted(per_country_src.glob(f"*_{DATE}.md"))
    print(f"\n→ 转换 {len(md_files)} 份国家报告...")
    ok_docx = ok_pdf = 0
    for i, src in enumerate(md_files, 1):
        md_dst = OUT / "per_country" / "md" / src.name
        docx_dst = OUT / "per_country" / "docx" / f"{src.stem}.docx"
        pdf_dst = OUT / "per_country" / "pdf" / f"{src.stem}.pdf"
        shutil.copy2(src, md_dst)
        try:
            md_to_docx(src, docx_dst)
            ok_docx += 1
        except Exception as e:
            print(f"  ✗ {src.stem} docx: {e}")
        try:
            md_to_pdf(src, pdf_dst)
            ok_pdf += 1
        except Exception as e:
            print(f"  ✗ {src.stem} pdf: {e}")
        if i % 10 == 0:
            print(f"  [{i}/{len(md_files)}] done (docx={ok_docx}, pdf={ok_pdf})")

    print(f"\n✓ per_country: {len(md_files)} md / {ok_docx} docx / {ok_pdf} pdf")

    # 5. 根目录索引 README
    readme = OUT / "README.md"
    readme_content = f"""# BRI 64 国中文教育指标体系 · 2026-04-24 版最终交付包

## 目录结构

- `reports/`              — 4 份核心报告（md + docx + pdf）
  - 分析报告_对比分析（Top20 / 最难指标 / 置信度等）
  - 项目状态总览
  - 可疑格详细清单（447 格）
  - 精修回写摘要
- `xlsx/`                 — Excel 数据
  - `bri_64country_merged_{DATE}.xlsx` — **合并总表：64 国 × 68 指标**
  - `*_refined447_{DATE}.xlsx` (13 份)  — 各批次精修版
  - `bri_421_suspicious_targets_{DATE}.json` — 可疑格目标
  - `bri_refine447_ckpt.json` — 精修 447 完整推导 + 源 URL
- `charts/`               — 6 张对比图 PNG
- `per_country/`          — 64 国独立指标详报
  - `md/`   MD 源文件
  - `docx/` Word 版
  - `pdf/`  PDF 版
  每份含该国 68 指标逐行详情 + 计算依据 + 源 URL + 精修 derivation

## 总体数据质量

- **总格**：4,352（64 国 × 68 指标）
- **扎实率**：86.1%（数字 46.0% + 定性 40.1%）
- **弱数据**：13.9%（0 / 无 / 有限 等）
- **最扎实国**：新加坡（92.8%）
- **最薄弱国**：不丹（73.9%）

生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}
"""
    readme.write_text(readme_content, encoding="utf-8")
    print(f"\n✓ README: {readme}")
    print(f"\n全部完成。主目录：{OUT}")


if __name__ == "__main__":
    main()
