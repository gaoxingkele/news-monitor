from __future__ import annotations
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

OUT = Path(r'D:\BaiduSyncdisk\aicoding\news-monitor\topics')
OUT.mkdir(parents=True, exist_ok=True)
DATE = '20260410'
pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
styles = getSampleStyleSheet()
TITLE = ParagraphStyle('T', parent=styles['Title'], fontName='STSong-Light', fontSize=18, leading=24, spaceAfter=12)
BODY = ParagraphStyle('B', parent=styles['BodyText'], fontName='STSong-Light', fontSize=11, leading=16, spaceAfter=6)
HEAD = ParagraphStyle('H', parent=styles['Heading1'], fontName='STSong-Light', fontSize=14, leading=18, spaceBefore=8, spaceAfter=8)

DATA = {
    'malaysia': {
        'cn':'马来西亚','title':'马来西亚中文教育发展态势评估',
        'summary':['马来西亚中文教育呈现国家学校体系、社群办学体系和高校/研究体系并行的多轨格局。','中文不具官方语言地位，但在华小（SJKC）层面拥有明确教学媒介地位。','社会基础极强，但多源流教育长期伴随政策争议。'],
        'rows':[
            (1,'外交伙伴关系等级','已确认','全面战略伙伴关系并持续深化','https://www.kln.gov.my/web/guest/home?_101_assetEntryId=9861626&_101_struts_action=%2Fasset_publisher%2Fview_content&_101_type=content&_101_urlTitle=joint-statement-between-the-people-s-republic-of-china-and-malaysia-on-deepening-the-comprehensive-strategic-partnership-towards-china-malaysia-commun&inheritRedirect=false&p_p_id=101&p_p_lifecycle=0&p_p_mode=view&p_p_state=maximized'),
            (2,'纳入教育体系层级','已确认','基础教育明确纳入；中学、高教、社区教育持续存在','https://www.moe.gov.my/storage/files/shares/Dasar/Dasar%20Pendidikan%20Kebangsaan/Dasar%20Pendidikan%20Kebangsaan%20Edisi%20Keempat.pdf'),
            (3,'中文教育地位','已确认','非官方语言；在华小中为主要教学媒介','https://www.malaysia.gov.my/en/government/kenali-malaysia/bahasa-rasmi'),
            (10,'教育政策环境风险','部分确认','中等','https://www.moe.gov.my/sistem-pendidikan'),
            (17,'研究机构/智库','已确认','有','https://fcs.utar.edu.my/'),
            (37,'硕博培养点','已确认','有','https://study.um.edu.my/master-of-chinese-studies-coursework'),
            (42,'开设中文专业高校数','已确认','至少2所','https://study.utar.edu.my/chinese-studies.php'),
            (45,'孔子学院/中文工坊','已确认','至少1个','https://www.um.edu.my/global-engagement'),
            (48,'教师/教育协会','已确认','有且网络化程度高','https://www.dongzong.my/en/about-us/'),
        ],
        'conclusion':'社会基础最强，但政策争议长期存在。'
    },
    'thailand': {
        'cn':'泰国','title':'泰国中文教育发展态势评估',
        'summary':['泰国中文教育属于需求驱动强、政策开放度较高的外语教育扩张型样本。','中文整体上仍是高热度外语，而非制度性母语或官方语言。','大学体系和华社机构支撑明显，是其增长的重要基础。'],
        'rows':[
            (1,'外交伙伴关系等级','已确认','全面战略合作伙伴关系','https://mfa.go.th/en/content/cscp-2024-en?cate=5d5bcb4e15e39c306000683e'),
            (2,'纳入教育体系层级','部分确认','学校项目与高校体系已存在','https://www.moe.go.th/%E0%B8%AB%E0%B9%89%E0%B8%AD%E0%B8%87%E0%B9%80%E0%B8%A3%E0%B8%B5%E0%B8%A2%E0%B8%99%E0%B8%A0%E0%B8%B2%E0%B8%A9%E0%B8%B2%E0%B8%88%E0%B8%B5%E0%B8%99/'),
            (3,'中文教育地位','已确认','高热度外语；非官方语言','https://thailand.go.th/page/thai-language'),
            (10,'教育政策环境风险','部分确认','低到中','https://www.moe.go.th/%E0%B8%AB%E0%B9%89%E0%B8%AD%E0%B8%87%E0%B9%80%E0%B8%A3%E0%B8%B5%E0%B8%A2%E0%B8%99%E0%B8%A0%E0%B8%B2%E0%B8%A9%E0%B8%B2%E0%B8%88%E0%B8%B5%E0%B8%99/'),
            (17,'研究机构/智库','已确认','有','https://hcuchinamonitor.hcu.ac.th/en/home-en'),
            (37,'硕博培养点','已确认','有','https://www.arts.chula.ac.th/east/chinese/Curriculum.html'),
            (39,'本土教师培养项目','部分确认','存在','https://www.arts.chula.ac.th/east/chinese/history_en.html'),
            (42,'开设中文专业高校数','已确认','至少1所成熟样本','https://www.arts.chula.ac.th/east/chinese/index_en.html'),
            (63,'中文+职业技能','部分确认','存在明显商务导向','https://hcuchinamonitor.hcu.ac.th/en/home-en'),
        ],
        'conclusion':'需求旺盛、扩张快，但中文仍属外语赛道。'
    },
    'indonesia': {
        'cn':'印尼','title':'印度尼西亚中文教育发展态势评估',
        'summary':['印尼中文教育处于恢复后扩张、政策逐步正规化的成长阶段。','国家课程体系已明确设置 Bahasa Mandarin 学习内容与评测标准。','整体模式更像“国家允许+市场拉动+高校和孔院支撑”的混合型。'],
        'rows':[
            (1,'外交伙伴关系等级','已确认','全面战略伙伴关系','https://www.mfa.gov.cn/eng/zy/jj/2022/cxesgjtytjhtg/202211/t20221117_10976784.html'),
            (2,'纳入教育体系层级','已确认','中学课程标准与评测体系已进入；高校体系存在','https://guru.kemdikbud.go.id/kurikulum/referensi-penerapan/capaian-pembelajaran/sd-sma/bahasa-mandarin/'),
            (3,'中文教育地位','已确认','外语科目；非官方语言','https://indonesia.go.id//kategori/sosial-budaya/10256/bahasa-indonesia-resmi-digunakan-sebagai-bahasa-kerja-di-sidang-umum-unesco?lang=1%3Flang%3D1'),
            (10,'教育政策环境风险','部分确认','中等','https://guru.kemdikbud.go.id/kurikulum/referensi-penerapan/capaian-pembelajaran/sd-sma/bahasa-mandarin/'),
            (17,'研究机构/智库','部分确认','有高校与孔院支撑','https://international.ui.ac.id/humanities-bachelors/'),
            (39,'本土教师培养项目','部分确认','存在教师认证/专业化路径','https://pusatinformasi.guru.kemdikbud.go.id/hc/en-us/articles/35466120859161-Informasi-Tentang-Rumpun-Bidang-Studi-Pada-Sertifikasi-Pendidik'),
            (42,'开设中文专业高校数','已确认','至少2所','https://chinese.binus.ac.id/'),
            (45,'孔子学院/中文工坊','已确认','至少1个','https://jurnal.uns.ac.id/maobi/about/contact'),
            (64,'教师资格认证','已确认','有','https://pusatinformasi.guru.kemdikbud.go.id/hc/en-us/articles/35466120859161-Informasi-Tentang-Rumpun-Bidang-Studi-Pada-Sertifikasi-Pendidik'),
        ],
        'conclusion':'国家课程已承认，但整体制度深度仍在上升。'
    },
    'vietnam': {
        'cn':'越南','title':'越南中文教育发展态势评估',
        'summary':['越南中文教育的突出特点是课程承认较明确、高校与研究机构较强、教师培训活跃。','中文已进入普通教育外语1课程框架，高校和研究机构支撑明显。','挑战主要来自英语、日语、韩语等竞争语种，而非行政层面的直接打压。'],
        'rows':[
            (1,'外交伙伴关系等级','已确认','全面战略合作伙伴关系并推进命运共同体建设','https://mofa.gov.vn/web/ministry-of-foreign-affairs/detail/chi-tiet/developing-ties-with-china-a-strategic-priority-in-vietnam-s-foreign-policy-deputy-pm-58405-591.html'),
            (2,'纳入教育体系层级','已确认','中学阶段已进入外语1课程框架；高校体系成熟','https://moet.gov.vn/giaoducquocdan/day-va-hoc-ngoai-ngu/Pages/Default.aspx?ItemID=7461'),
            (3,'中文教育地位','已确认','正式外语课程选项；非国家共同语言','https://mofa.gov.vn/web/guest/dan-toc-ngon-ngu'),
            (10,'教育政策环境风险','部分确认','中低','https://moet.gov.vn/giaoducquocdan/day-va-hoc-ngoai-ngu/Pages/Default.aspx?ItemID=7461'),
            (17,'研究机构/智库','已确认','有','https://ussh.vnu.edu.vn/en/organizational-structure/centers-institutes-companies/center-for-chinese-studies-13067.html'),
            (37,'硕博培养点','已确认','有','https://ulis.vnu.edu.vn/co-cau-to-chuc/khoa-nnvh-trung/'),
            (39,'本土教师培养项目','已确认','有','https://ulis.vnu.edu.vn/khai-mac-khoa-boi-duong-giao-vien-tieng-trung-tai-viet-nam/'),
            (42,'开设中文专业高校数','已确认','至少2所','https://chinese.ftu.edu.vn/index.php/gi%E1%BB%9Bi-thi%E1%BB%87u/l%E1%BB%8Bch-s%E1%BB%AD-h%C3%ACnh-th%C3%A0nh'),
            (45,'孔子学院/中文工坊','已确认','至少1个','https://en.ulis.vnu.edu.vn/blog/archives/ulis-students-won-prizes-at-20th-chinese-bridge-competition/'),
            (64,'教师资格认证','部分确认','存在较清晰培养与认证路径','https://en.ulis.vnu.edu.vn/aboutulis-organizational-structure/faculty-of-chinese-language-and-culture/ulis-sunwah-chinese-teaching-and-research-centre-ulis-sunwah/'),
        ],
        'conclusion':'课程承认明显，大学与教师培训支撑较突出。'
    },
}

COMPARE = [
    ('新加坡','高制度化样本','很高','高','很强','低','最适合作为标准样例。'),
    ('马来西亚','社会基础最强样本','高','很高','很强','中','社会组织化强，但政策争议长期存在。'),
    ('泰国','外语扩张型样本','中高','中高','强','低到中','需求旺盛、扩张快，但中文仍属外语赛道。'),
    ('印尼','成长型样本','中','中','中','中','国家课程已承认，但整体制度深度仍在上升。'),
    ('越南','增长型样本','中高','中','强','中低','课程承认明显，大学与教师培训支撑较突出。'),
]

def save_country(key, item):
    base = f'bri_{key}_assessment_{DATE}'
    md = OUT / f'{base}.md'
    docx = OUT / f'{base}.docx'
    pdf = OUT / f'{base}.pdf'
    xlsx = OUT / f'{base}_data.xlsx'

    lines = [f"# {item['title']}", '', f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}", '', '## 摘要', '']
    lines.extend([f'- {x}' for x in item['summary']])
    lines.extend(['', '## 核心汇总数据', ''])
    for r in item['rows']:
        lines.append(f'- 指标{r[0]} {r[1]}：{r[3]}（{r[2]}）')
        lines.append(f'  来源：{r[4]}')
    lines.extend(['', '## 结论', '', f"- {item['conclusion']}"])
    md.write_text('\n'.join(lines), encoding='utf-8')

    d = Document()
    for s in d.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2); s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
    h = d.add_heading(item['title'], level=0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = d.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_heading('摘要', level=1)
    for x in item['summary']: d.add_paragraph(x)
    d.add_heading('核心汇总数据', level=1)
    t = d.add_table(rows=1, cols=5)
    for i, h2 in enumerate(['指标ID','字段','状态','当前值','来源']): t.rows[0].cells[i].text = h2
    for r in item['rows']:
        row = t.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text, row[4].text = str(r[0]), r[1], r[2], r[3], r[4]
    d.add_heading('结论', level=1); d.add_paragraph(item['conclusion'])
    d.save(docx)

    story = [Paragraph(item['title'], TITLE), Paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}", BODY), Spacer(1,0.2*cm), Paragraph('摘要', HEAD)]
    for x in item['summary']: story.append(Paragraph(x, BODY))
    story.append(Paragraph('核心汇总数据', HEAD))
    for r in item['rows']:
        story.append(Paragraph(f'<b>指标{r[0]} {r[1]}：</b>{r[3]}（{r[2]}）<br/><font color="blue">{r[4]}</font>', BODY))
    story.append(Paragraph('结论', HEAD)); story.append(Paragraph(item['conclusion'], BODY))
    SimpleDocTemplate(str(pdf), pagesize=A4, leftMargin=2.2*cm, rightMargin=2.2*cm, topMargin=2*cm, bottomMargin=2*cm).build(story)

    wb = Workbook(); ws = wb.active; ws.title = 'core_data'
    ws.append(['metric_id','field','status','value','source'])
    fill = PatternFill('solid', fgColor='D9EAF7')
    for c in ws[1]: c.font = Font(bold=True); c.fill = fill; c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
    for r in item['rows']: ws.append(list(r))
    for row in ws.iter_rows(min_row=2):
        for c in row: c.alignment = Alignment(vertical='top', wrap_text=True)
    for col, width in {'A':10,'B':24,'C':12,'D':34,'E':78}.items(): ws.column_dimensions[col].width = width
    wb.save(xlsx)
    return [md, docx, pdf, xlsx]

all_paths = []
for k, v in DATA.items():
    all_paths.extend(save_country(k, v))

base = f'bri_5country_comparison_{DATE}'
md = OUT / f'{base}.md'; docx = OUT / f'{base}.docx'; pdf = OUT / f'{base}.pdf'; xlsx = OUT / f'{base}.xlsx'
wb = Workbook(); ws = wb.active; ws.title = 'comparison'
ws.append(['国家','总体定位','制度化程度','社会基础','研究与师资','风险判断','一句话结论'])
fill = PatternFill('solid', fgColor='D9EAF7')
for c in ws[1]: c.font = Font(bold=True); c.fill = fill; c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
for r in COMPARE: ws.append(list(r))
for row in ws.iter_rows(min_row=2):
    for c in row: c.alignment = Alignment(vertical='top', wrap_text=True)
for col, width in {'A':12,'B':18,'C':12,'D':12,'E':12,'F':10,'G':42}.items(): ws.column_dimensions[col].width = width
wb.save(xlsx)
md.write_text('\n'.join(['# 五国中文教育态势横向对比','',f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}",'','## 对比结论',''] + [f'- {r[0]}：{r[6]}' for r in COMPARE]), encoding='utf-8')
d = Document(); h = d.add_heading('五国中文教育态势横向对比', level=0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in COMPARE: d.add_paragraph(f'{r[0]}：{r[6]}')
d.save(docx)
story = [Paragraph('五国中文教育态势横向对比', TITLE), Spacer(1,0.2*cm)]
for r in COMPARE: story.append(Paragraph(f'<b>{r[0]}</b>：{r[6]}', BODY))
SimpleDocTemplate(str(pdf), pagesize=A4, leftMargin=2.2*cm, rightMargin=2.2*cm, topMargin=2*cm, bottomMargin=2*cm).build(story)
all_paths.extend([md, docx, pdf, xlsx])
for p in all_paths: print(p)
