from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle


TITLE = "“台方安全局在郑习会期间发布业务报告”舆情分类汇总与评估"
SUBTITLE = "时间窗口：2026-04-01 至 2026-04-07"

SECTIONS = [
    {
        "title": "一、总体判断",
        "paragraphs": [
            "本轮舆情的核心，不是单纯围绕台“国安局”业务报告文本本身展开，而是围绕“谁来定义郑习会”发生了明显的叙事争夺。绿营及安全系统试图将郑习会重构为“大陆威胁升级”“统战与认知战延伸”的安全事件；蓝营则将其包装为“和平沟通、危机降温、争取对话”的政治动作。",
            "从热度看，台湾岛内热度最高，且网页媒体、社交平台和论坛均同步放大；美国、日本、新加坡/东南亚媒体关注度次之，重点落在台海稳定、军购预算、北京对台影响力以及美中台三方信号博弈；欧洲和韩国关注较弱。",
            "从情绪看，社交媒体与论坛层面的警戒和批评情绪显著高于传统媒体；而国际主流媒体多采取相对克制的“风险评估”或“背景解释”框架，并未完全跟随岛内最强烈的威胁叙事。"
        ],
    },
    {
        "title": "二、分类汇总",
        "paragraphs": [
            "1. 安全威胁叙事：台安全系统、民进党与部分英文台媒将报告与郑习会绑定，强调“压拉两手策略”“渗透窃密”“灰区侵扰”“认知战”等标签，并将其外延到军购、网络安全、选举安全和国际情报合作。",
            "2. 和平交流/降温叙事：蓝营及部分中性媒体将郑习会定义为降低台海风险、争取沟通空间的举措，强调“对等、尊严、和平”，以对冲绿营的安全焦虑叙事。",
            "3. 军购与吓阻能力争议：美国相关报道以及岛内陆委会、英文媒体持续把郑习会与对美军购、特别防务预算和对外安全合作挂钩，构成国际关注的焦点之一。",
            "4. 统战/认知战指控：论坛和社交媒体最容易放大这一叙事，常把郑习会直接等同于“统战样板”“卖台”“替北京释放政治信号”。",
            "5. 国际信号博弈：新加坡、日本、澳大利亚等区域媒体最关注此事在北京、华盛顿、台北三边关系中的信号意义，而不只把它看成单纯的两岸会面。"
        ],
    },
    {
        "title": "三、平台与地区热度判断",
        "paragraphs": [
            "台湾岛内：高。网页媒体、政党表态、论坛与社交平台同步高热，且蓝绿对冲明显。",
            "美国：高。重点不是台“国安局”报告本身，而是其与台美军购、吓阻能力和对华警惕的关联。",
            "日本：中高。重点放在台海稳定、区域危机管理与郑习会的外溢安全影响。",
            "新加坡/东南亚：中高。主视角是北京—华盛顿—台北三方信号博弈与区域稳定。",
            "澳大利亚：中。聚焦和解路线与联盟压力并存、北京对台湾政局影响。",
            "欧洲：低到中。报道多为背景性和解释性，直接高热舆论较少。",
            "韩国：低。公开连续讨论不多，热度弱于日美与岛内。",
            "Reddit：中。英文圈与台湾小圈层有讨论，但热度不及 X 与台湾论坛。",
            "PTT/Disp：高。情绪最强，选战化和标签化最明显。",
            "Mobile01：中。存在较多反绿、反“国安操作”声音，讨论结构与 PTT 不同。",
            "Facebook：证据较弱。公开搜索索引可见度低，暂不宜作为高热主证据源。"
        ],
    },
    {
        "title": "四、论坛层进一步深挖",
        "paragraphs": [
            "Dcard：近期可核验讨论明显增加，讨论形态包括新闻转贴串、立场评论串以及投票帖。整体上，批评郑习会、质疑其为中共主导安排、将其与“国共一家”“卖台/统战”绑定的声音更强；但也能看到少量“和平交流总比对抗好”的中性或支持性表达。",
            "巴哈姆特：公开搜索索引中未稳定检出与本事件高度相关、且讨论量足够的近期哈啦板主串，因此不宜强行作为高代表性证据来源。现阶段可以判断巴哈并非本事件的主要舆情承载平台。",
            "LIHKG：在本轮时间窗口内未找到同等热度的新串，但可检出与郑丽文“中国人论述”直接相关的旧讨论。它更像是对郑个人路线与身份认同的长期负面背景情绪，而非这次事件的当期主战场。",
            "英文论坛：除Reddit既有样本外，新增可核验的 r/worldnews 讨论串热度较高，主叙事集中在‘谈比打好’、‘KMT是否更亲北京’、‘美国是否仍会坚定支持台湾’三条线上。相较之下，其他英文论坛对本事件本身的直接讨论较弱，更多停留在泛台海议题。"
        ],
    },
]

EVIDENCE = [
    {
        "category": "网页媒体/岛内",
        "items": [
            ("联合报", "将郑习会与国安报告直接并置，强调大陆‘压拉两手策略’。", "https://udn.com/news/story/124802/9424745"),
            ("Taipei Times", "将郑习会与阻挠军购预算、影响台美合作相联系。", "https://www.taipeitimes.com/News/taiwan/archives/2026/04/03/2003854954"),
            ("公视", "蓝营强调和平与对等尊严，绿营要求停止军压与打压。", "https://news.pts.org.tw/article/802322"),
            ("Taiwan News", "英文舆论框架下突出高科技人才、渗透与安全威胁。", "https://www.taiwannews.com.tw/en/news/6335148"),
        ],
    },
    {
        "category": "国际媒体",
        "items": [
            ("AP", "美国议员公开支持台湾特别防务预算，形成对郑习会的外部对照。", "https://apnews.com/article/685b8cf5feef733a86b360325913e442"),
            ("Japan Times", "关注郑习会是否影响区域危机管理和台海稳定。", "https://www.japantimes.co.jp/news/2026/04/01/asia-pacific/politics/taiwain-opposition-leader-china/"),
            ("CNA Singapore", "明确将其视为北京、华盛顿、台北三方信号博弈。", "https://www.channelnewsasia.com/east-asia/kmt-chair-china-trip-signals-beijing-washington-taipei-6036701"),
            ("DW", "欧洲视角更偏背景解释和政治含义判断。", "https://amp.dw.com/en/taiwan-opposition-leader-accepts-xis-invite-to-visit-china/a-76588423"),
        ],
    },
    {
        "category": "X/Twitter 大V与媒体账号",
        "items": [
            ("Reuters", "中性报道‘和平使命’，但评论区明显分裂：和平降温 vs 国民党被北京利用。", "https://x.com/i/status/2041009274789257314"),
            ("TaiwanMonitor", "观点最警戒，认为郑习会意在阻挠美台军售、内政化两岸问题。", "https://x.com/i/status/2040103245624791509"),
            ("Focus Taiwan", "转述陆委会口径，突出‘阻断军购’叙事。", "https://x.com/i/status/2039709436386615767"),
            ("Jojje Olsson", "偏怀疑，强调其非官方性质和象征意味。", "https://x.com/i/status/2040989499493343726"),
        ],
    },
    {
        "category": "Reddit / 论坛",
        "items": [
            ("Reddit r/China", "英文圈把焦点放在KMT是否进一步坐实亲北京路线，以及对台美防务合作的影响。", "https://www.reddit.com/r/China/comments/1s80s72/taiwan_opposition_leader_accepts_xis_invite_to/"),
            ("Reddit r/worldnews", "高热英文串里既有“谈比打好”的现实主义立场，也有将郑习会视为KMT更公开靠近北京的批评。", "https://www.reddit.com/r/worldnews/comments/1s7h3y1/taiwans_kuomintang_leader_to_visit_mainland_china/"),
            ("Reddit r/Taiwanese", "中文小圈层更倾向把郑习会视为统战样板和安全风险。", "https://www.reddit.com/r/Taiwanese/comments/1sdu842/%E9%84%AD%E9%BA%97%E6%96%87%E5%8D%B3%E5%B0%87%E8%A8%AA%E4%B8%AD%E4%BF%83%E5%92%8C%E5%B9%B3_%E5%88%86%E6%9E%90%E6%8C%87%E5%8C%97%E4%BA%AC%E6%88%96%E8%97%89%E6%A9%9F%E6%96%BD%E5%A3%93%E8%BB%8D%E5%94%AE/"),
            ("PTT HatePolitics", "即时政治情绪最强，讨论大量转向选战、加分减分与‘卖台/抹红’标签战。", "https://www.ptt.cc/bbs/HatePolitics/M.1774838712.A.083.html"),
            ("Mobile01", "可见较多反绿、反‘国安叙事操作’的逆向评论。", "https://www.mobile01.com/topicdetail.php?f=638&t=7244918"),
            ("Dcard 讨论串", "直接把‘郑习会’定义为中共主导的套装行程，评论区批评与质疑声量较高。", "https://www.dcard.tw/f/talk/p/261238817"),
            ("Dcard 投票串", "围绕‘郑是否应主动提中华民国’出现投票式分化，显示平台用户更偏身份认同与象征表达层面的争论。", "https://www.dcard.tw/f/trending/p/261208477"),
            ("Dcard 新闻转贴", "三立相关新闻在Dcard形成较高留言量，‘国共一家亲’与‘别配合统战演出’是显著口号。", "https://www.dcard.tw/%40setn54/post/261204689"),
            ("LIHKG 背景串", "虽非本轮时间窗口高热新串，但可见香港讨论区对郑丽文‘中国人论述’长期负面评价，构成外部华语圈背景情绪。", "https://lihkg.com/thread/3979309/page/1"),
        ],
    },
]

CONCLUSION = [
    "综合看，这起事件的舆情并非围绕“报告真假”单线展开，而是迅速演化为一场“安全 vs 和平、军购 vs 交流、统战指控 vs 反政治操作”的叙事争夺。",
    "岛内是主战场，国际媒体更多从台海稳定和三边信号角度观察，并未完全接受台安全系统的最高警报框架。",
    "社交媒体和论坛层面，警戒、批评和标签化情绪明显强于传统媒体；它们更适合作为情绪温度计，而不是政策事实本身的直接依据。",
    "在整体评估中，最有代表性的证据组合应当是：岛内网页媒体 + 国际主流媒体 + X大V与媒体账号 + 台湾论坛/Reddit/Dcard 讨论。Facebook目前公开可核验证据不足，建议仅作为补充观察对象。",
    "新增论坛层深挖显示：Dcard比预期更接近年轻用户的即时政治情绪池；Reddit英文圈更偏‘台美防务与路线选择’；LIHKG对本案当期热度有限，但提供了郑个人标签化认知的背景参照。"
]


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_docx(output_path: Path) -> None:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    p = doc.add_heading(TITLE, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph(SUBTITLE)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3 = doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for section in SECTIONS:
        doc.add_heading(section["title"], level=1)
        for text in section["paragraphs"]:
            para = doc.add_paragraph(text)
            para.paragraph_format.space_after = Pt(6)

    doc.add_heading("四、代表性舆情观点摘要与链接", level=1)
    for group in EVIDENCE:
        doc.add_heading(group["category"], level=2)
        for name, summary, url in group["items"]:
            para = doc.add_paragraph(style=None)
            para.add_run(f"{name}：").bold = True
            para.add_run(summary + " 链接：")
            add_hyperlink(para, url, url)

    doc.add_heading("五、结论", level=1)
    for text in CONCLUSION:
        doc.add_paragraph(text)

    doc.save(output_path)


def build_pdf(output_path: Path) -> None:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCN",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=18,
        leading=24,
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "BodyCN",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=11,
        leading=16,
        spaceAfter=6,
    )
    heading_style = ParagraphStyle(
        "HeadingCN",
        parent=styles["Heading1"],
        fontName="STSong-Light",
        fontSize=14,
        leading=18,
        spaceBefore=8,
        spaceAfter=8,
    )
    subheading_style = ParagraphStyle(
        "SubHeadingCN",
        parent=styles["Heading2"],
        fontName="STSong-Light",
        fontSize=12,
        leading=16,
        spaceBefore=6,
        spaceAfter=6,
    )

    story = [
        Paragraph(TITLE, title_style),
        Paragraph(SUBTITLE, body_style),
        Paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}", body_style),
        Spacer(1, 0.3 * cm),
    ]

    for section in SECTIONS:
        story.append(Paragraph(section["title"], heading_style))
        for text in section["paragraphs"]:
            story.append(Paragraph(text, body_style))

    story.append(Paragraph("四、代表性舆情观点摘要与链接", heading_style))
    for group in EVIDENCE:
        story.append(Paragraph(group["category"], subheading_style))
        for name, summary, url in group["items"]:
            txt = f"<b>{name}：</b>{summary}<br/><font color='blue'>{url}</font>"
            story.append(Paragraph(txt, body_style))

    story.append(Paragraph("五、结论", heading_style))
    for text in CONCLUSION:
        story.append(Paragraph(text, body_style))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=2.2 * cm,
        rightMargin=2.2 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.0 * cm,
    )
    doc.build(story)


def main() -> None:
    out_dir = Path(r"D:\BaiduSyncdisk\2026-04\台情")
    out_dir.mkdir(parents=True, exist_ok=True)
    base = "台方安全局郑习会业务报告舆情分类汇总与评估_20260407_增强版"
    docx_path = out_dir / f"{base}.docx"
    pdf_path = out_dir / f"{base}.pdf"
    build_docx(docx_path)
    build_pdf(pdf_path)
    print(docx_path)
    print(pdf_path)


if __name__ == "__main__":
    main()
