"""Deep analysis: semantic clustering + 2000-word analytical report.

Reads from filtered JSON (output of fetch_event.py), calls Grok reasoning
for deep clustering and final report generation.

Usage:
    python analyze_event.py output/reports/XXXX_boao_final.json
    python analyze_event.py   # auto-find latest *_boao*.json
"""
from __future__ import annotations

import asyncio
import json
import os
import re
import sys
from datetime import datetime, timezone
from pathlib import Path

if sys.platform == "win32":
    os.environ.setdefault("PYTHONIOENCODING", "utf-8")
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "src"))

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

from news_monitor.proxy import get_proxies_for_url, build_httpx_client
from news_monitor.cloubic import resolve_openai_compatible_endpoint
from news_monitor.media_tier import get_tier, get_tier_label
from news_monitor.trend import record_snapshot


# ── Grok API helper ──────────────────────────────────────────────────────────

async def _call_grok(
    system_prompt: str,
    user_msg: str,
    overseas_proxy: str,
    model: str = "grok-4-1-fast-reasoning",
    temperature: float = 0.3,
    timeout: float = 180.0,
) -> str:
    """Single Grok API call. Returns response text."""
    api_key = os.environ.get("GROK_API_KEY", "")
    if not api_key:
        raise RuntimeError("GROK_API_KEY not set")

    api_key, base_url, model_chain, via_cloubic = resolve_openai_compatible_endpoint(
        "grok",
        direct_api_key=api_key,
        direct_base_url=os.environ.get("GROK_BASE_URL", "") or "https://api.x.ai/v1",
        direct_model=model,
        reasoning=True,
    )
    if not api_key:
        raise RuntimeError("resolved Grok/Cloubic API key not set")

    base_url = base_url.rstrip("/")
    if base_url.endswith("/chat/completions"):
        base_url = base_url[: -len("/chat/completions")]
    url = f"{base_url}/chat/completions"
    proxy = None if via_cloubic else get_proxies_for_url(url, overseas_proxy)

    body = {
        "model": model_chain[0],
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_msg},
        ],
        "temperature": temperature,
    }

    async with build_httpx_client(proxy_url=proxy or "", timeout=timeout) as client:
        resp = await client.post(url, json=body, headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        })
        resp.raise_for_status()
        data = resp.json()

    content = data["choices"][0]["message"]["content"].strip()
    return content


def _repair_json(text: str) -> list | dict | None:
    """Try to parse JSON from LLM response."""
    text = text.strip()
    # Strip markdown fences
    if text.startswith("```"):
        text = text.split("\n", 1)[-1].rsplit("```", 1)[0].strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Truncated array
    if text.startswith("["):
        last = text.rfind("}")
        if last > 0:
            try:
                return json.loads(text[:last + 1].rstrip(", \n\t") + "]")
            except json.JSONDecodeError:
                pass

    # Truncated object
    if text.startswith("{"):
        last = text.rfind("}")
        if last > 0:
            try:
                return json.loads(text[:last + 1])
            except json.JSONDecodeError:
                pass

    return None


# ── Step 1: Semantic clustering ──────────────────────────────────────────────

_CLUSTER_PROMPT = """\
你是一位资深国际舆情分析师。下面是关于"博鳌亚洲论坛2026"的海外媒体报道列表。

请将这些文章按照**实际讨论的议题/事件**进行语义聚类，合并为 15-25 个主题群。

聚类规则：
1. 同一事件的不同媒体报道合并为一个主题
2. 相似议题（如"AI讨论"和"科技创新"）可合并
3. 每个主题必须有明确的主题名（中文，15字以内）
4. 标注每个主题下的舆情态度倾向（正面/中性/负面/争议）
5. 提取该主题下最有代表性的 2-3 个核心观点

返回 JSON 数组：
[
  {
    "topic": "主题名称",
    "sentiment": "positive|neutral|negative|mixed",
    "article_ids": [0, 3, 15, ...],
    "key_points": ["核心观点1", "核心观点2"],
    "importance": 1-5
  },
  ...
]

只返回 JSON，不要加任何解释。
"""


async def _semantic_cluster(
    articles: list[dict],
    overseas_proxy: str,
) -> list[dict]:
    """Send all articles to Grok for semantic clustering. May split into batches."""
    # Build compact input: id + title + sentiment
    items = []
    for i, a in enumerate(articles):
        title = a.get("title_zh") or a.get("title", "")
        cat = a.get("category", "neutral")
        items.append({"id": i, "title": title[:80], "s": cat})

    # Estimate: ~100 chars per article input, ~200 tokens per cluster output
    # 438 articles ≈ 44KB input → fits in 1 call
    user_msg = json.dumps(items, ensure_ascii=False)
    input_size = len(user_msg)
    print(f"  [clustering] {len(items)} articles, input {input_size//1024}KB", flush=True)

    # If too large (>500KB), split into 2 calls and merge
    if input_size > 500_000:
        mid = len(items) // 2
        print(f"  [clustering] splitting into 2 batches of {mid}...", flush=True)
        msg1 = json.dumps(items[:mid], ensure_ascii=False)
        msg2 = json.dumps(items[mid:], ensure_ascii=False)

        r1 = await _call_grok(_CLUSTER_PROMPT, msg1, overseas_proxy, timeout=240.0)
        c1 = _repair_json(r1) or []
        print(f"  [clustering] batch 1: {len(c1)} clusters", flush=True)

        r2 = await _call_grok(_CLUSTER_PROMPT, msg2, overseas_proxy, timeout=240.0)
        c2 = _repair_json(r2) or []
        # Offset IDs in batch 2
        for c in c2:
            c["article_ids"] = [x + mid for x in c.get("article_ids", [])]
        print(f"  [clustering] batch 2: {len(c2)} clusters", flush=True)

        return c1 + c2
    else:
        raw = await _call_grok(_CLUSTER_PROMPT, user_msg, overseas_proxy, timeout=240.0)
        clusters = _repair_json(raw) or []
        print(f"  [clustering] got {len(clusters)} clusters", flush=True)
        return clusters


# ── Step 2: Generate 2000-word analytical report ─────────────────────────────

_REPORT_PROMPT = """\
你是一位为中国政府外事部门撰写舆情简报的资深分析师。

基于以下关于"博鳌亚洲论坛2026"的海外媒体舆情聚类分析结果，撰写一份约2000字的**深度舆情分析报告**。

报告结构要求：
1. **舆情总览**（200字）：总体态势、报道热度、主要关注方向
2. **核心议题分析**（800字）：按重要度逐一分析排名前 8-10 的议题群，每个议题包含：
   - 海外媒体的主要观点和态度
   - 正面与负面声音的对比
   - 值得注意的具体论述或评论
3. **舆情风险与机遇**（500字）：
   - 负面舆情的主要风险点和传播趋势
   - 正面舆情可以利用的机遇
   - 需要回应或澄清的误解
4. **研判与建议**（500字）：
   - 短期舆情走向预判
   - 对外传播建议（哪些叙事有效、哪些需要调整）
   - 需重点关注的媒体和议题

写作风格：专业、客观、信息密度高、判断有据。避免空话套话。
直接输出报告全文，不要加"以下是报告"之类的引导语。
"""


async def _generate_deep_report(
    clusters: list[dict],
    articles: list[dict],
    dist: dict[str, int],
    overseas_proxy: str,
) -> str:
    """Generate 2000-word analytical report from clusters."""
    # Build input: cluster summaries + stats
    input_lines = [
        f"报道总量：{len(articles)} 篇海外媒体报道（已排除中国大陆媒体）",
        f"态度分布：正面{dist.get('positive',0)}篇 / 中性{dist.get('neutral',0)}篇 / 负面{dist.get('negative',0)}篇 / 分析{dist.get('analysis',0)}篇",
        "",
        "聚类结果（按重要度排序）：",
        "",
    ]

    # Sort clusters by importance
    clusters_sorted = sorted(clusters, key=lambda c: c.get("importance", 3), reverse=True)
    for i, c in enumerate(clusters_sorted, 1):
        topic = c.get("topic", "未命名")
        sentiment = c.get("sentiment", "neutral")
        n = len(c.get("article_ids", []))
        importance = c.get("importance", 3)
        points = c.get("key_points", [])

        sent_label = {"positive": "正面", "neutral": "中性", "negative": "负面", "mixed": "争议"}.get(sentiment, sentiment)
        input_lines.append(f"{i}. 【{topic}】（{n}篇，{sent_label}，重要度{importance}）")
        for p in points:
            input_lines.append(f"   - {p}")

        # Add representative article titles with tier info
        aids = c.get("article_ids", [])[:5]
        for aid in aids:
            if aid < len(articles):
                t = articles[aid].get("title_zh") or articles[aid].get("title", "")
                src = articles[aid].get("source_name", "")
                url = articles[aid].get("source_url", "")
                tier = get_tier(url) if url else 4
                tier_tag = f"T{tier}" if tier <= 3 else ""
                tier_str = f" ({tier_tag})" if tier_tag else ""
                input_lines.append(f"   [{src}{tier_str}] {t[:70]}")
        input_lines.append("")

    user_msg = "\n".join(input_lines)
    print(f"  [report] input {len(user_msg)//1024}KB, generating ~2000 words...", flush=True)

    raw = await _call_grok(_REPORT_PROMPT, user_msg, overseas_proxy, timeout=300.0)

    # Strip markdown fences
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[-1].rsplit("```", 1)[0].strip()

    print(f"  [report] generated {len(raw)} chars", flush=True)
    return raw


# ── Output ────────────────────────────────────────────────────────────────────

def _build_full_md(
    report_text: str,
    clusters: list[dict],
    articles: list[dict],
    dist: dict[str, int],
) -> str:
    """Build the final MD combining deep report + cluster details."""
    gen_time = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    lines = [
        "# 博鳌亚洲论坛2026 — 海外舆情深度分析报告",
        "",
        f"**生成时间：** {gen_time}  |  **数据来源：** Brave / Gemini / Grok（海外媒体）",
        "",
        f"> 基于 {len(articles)} 篇海外媒体报道的语义聚类与深度分析",
        f"> 正面 {dist.get('positive',0)} / 中性 {dist.get('neutral',0)} / 负面 {dist.get('negative',0)} / 分析 {dist.get('analysis',0)}",
        "",
        "---",
        "",
        report_text,
        "",
        "---",
        "",
        "# 附录：主题聚类详情",
        "",
    ]

    clusters_sorted = sorted(clusters, key=lambda c: c.get("importance", 3), reverse=True)
    sent_emoji = {"positive": "🟢", "neutral": "🔵", "negative": "🔴", "mixed": "🟡"}

    for i, c in enumerate(clusters_sorted, 1):
        topic = c.get("topic", "未命名")
        sentiment = c.get("sentiment", "neutral")
        n = len(c.get("article_ids", []))
        importance = c.get("importance", 3)
        points = c.get("key_points", [])
        emoji = sent_emoji.get(sentiment, "⚪")

        lines.append(f"## {i}. {emoji} {topic}（{n} 篇，{'★' * importance}）")
        lines.append("")

        if points:
            for p in points:
                lines.append(f"- {p}")
            lines.append("")

        # Article list
        aids = c.get("article_ids", [])
        for aid in aids[:10]:
            if aid < len(articles):
                a = articles[aid]
                t = a.get("title_zh") or a.get("title", "")
                src = a.get("source_name", "")
                url = a.get("source_url", "")
                if url:
                    lines.append(f"  - [{src}] {t[:70]} — {url}")
                else:
                    lines.append(f"  - [{src}] {t[:70]}")
        if len(aids) > 10:
            lines.append(f"  - *...及其他 {len(aids) - 10} 篇*")

        lines += ["", "---", ""]

    return "\n".join(lines)


def _build_docx(
    report_text: str,
    clusters: list[dict],
    articles: list[dict],
    dist: dict[str, int],
    output_path: str,
):
    """Build Word document with the deep analysis report."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    t = doc.add_heading("博鳌亚洲论坛2026 — 海外舆情深度分析报告", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    gen_time = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"生成时间：{gen_time}  |  基于 {len(articles)} 篇海外媒体报道")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(128, 128, 128)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"正面 {dist.get('positive',0)} / 中性 {dist.get('neutral',0)} / 负面 {dist.get('negative',0)} / 分析 {dist.get('analysis',0)}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()

    # Main report — split by section headers
    for para_text in report_text.split("\n"):
        para_text = para_text.strip()
        if not para_text:
            doc.add_paragraph()
            continue

        # Detect section headers (Chinese numbered headers or markdown ##)
        if para_text.startswith("## ") or para_text.startswith("# "):
            heading_text = para_text.lstrip("# ").strip()
            doc.add_heading(heading_text, level=1)
        elif para_text.startswith("### "):
            heading_text = para_text.lstrip("# ").strip()
            doc.add_heading(heading_text, level=2)
        elif re.match(r'^[一二三四五六七八九十]+[、.]', para_text):
            doc.add_heading(para_text, level=1)
        elif re.match(r'^\d+[、.\s]', para_text) and len(para_text) < 50:
            doc.add_heading(para_text, level=2)
        elif para_text.startswith("- ") or para_text.startswith("• "):
            doc.add_paragraph(para_text[2:], style='List Bullet')
        else:
            doc.add_paragraph(para_text)

    # Appendix: cluster details
    doc.add_page_break()
    doc.add_heading("附录：主题聚类详情", level=0)

    clusters_sorted = sorted(clusters, key=lambda c: c.get("importance", 3), reverse=True)
    sent_labels = {"positive": "正面", "neutral": "中性", "negative": "负面", "mixed": "争议"}
    sent_colors = {"positive": (34, 139, 34), "neutral": (70, 130, 180), "negative": (205, 92, 92), "mixed": (200, 150, 50)}

    for i, c in enumerate(clusters_sorted, 1):
        topic = c.get("topic", "未命名")
        sentiment = c.get("sentiment", "neutral")
        n = len(c.get("article_ids", []))
        importance = c.get("importance", 3)
        points = c.get("key_points", [])
        color = sent_colors.get(sentiment, (0, 0, 0))

        h = doc.add_heading(f"{i}. {topic}（{n}篇，{'★' * importance}）", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)

        if points:
            for pt in points:
                doc.add_paragraph(pt, style='List Bullet')

        aids = c.get("article_ids", [])[:8]
        for aid in aids:
            if aid < len(articles):
                a = articles[aid]
                t_text = a.get("title_zh") or a.get("title", "")
                src = a.get("source_name", "")
                p = doc.add_paragraph()
                r = p.add_run(f"[{src}] {t_text[:70]}")
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(output_path)


# ── Main ──────────────────────────────────────────────────────────────────────

async def main():
    from news_monitor.config_loader import load_config
    config = load_config("config.yaml")
    overseas_proxy = config.get("proxy", {}).get("overseas_proxy", "")

    # Find input JSON
    if len(sys.argv) > 1:
        json_path = sys.argv[1]
    else:
        # Auto-find latest boao JSON
        candidates = sorted(Path("output/reports").glob("*boao*.json"), reverse=True)
        # Prefer *_final.json
        finals = [c for c in candidates if "final" in c.stem]
        json_path = str(finals[0] if finals else candidates[0])

    print(f"Input: {json_path}")
    with open(json_path, "r", encoding="utf-8") as f:
        articles = json.load(f)
    print(f"Articles: {len(articles)}")

    dist = {}
    for a in articles:
        cat = a.get("category", "neutral")
        dist[cat] = dist.get(cat, 0) + 1
    print(f"Distribution: {dist}")

    # Step 1: Semantic clustering (1 API call)
    print("\n=== Step 1: Semantic clustering ===")
    clusters = await _semantic_cluster(articles, overseas_proxy)
    if not clusters:
        print("ERROR: clustering failed")
        return

    # Save clusters
    clusters_path = json_path.replace(".json", "_clusters.json")
    with open(clusters_path, "w", encoding="utf-8") as f:
        json.dump(clusters, f, ensure_ascii=False, indent=2)
    print(f"Clusters saved: {clusters_path}")

    # Step 2: Generate 2000-word report (1 API call)
    print("\n=== Step 2: Deep analysis report ===")
    report_text = await _generate_deep_report(clusters, articles, dist, overseas_proxy)
    if not report_text:
        print("ERROR: report generation failed")
        return

    # Step 3: Output
    print("\n=== Step 3: Output ===")
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    output_dir = "output/reports"

    # MD
    md_content = _build_full_md(report_text, clusters, articles, dist)
    md_path = f"{output_dir}/{ts}_boao_deep_analysis.md"
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"MD: {md_path}")

    # DOCX
    docx_path = f"{output_dir}/{ts}_boao_deep_analysis.docx"
    try:
        _build_docx(report_text, clusters, articles, dist, docx_path)
        print(f"DOCX: {docx_path}")
    except Exception as exc:
        print(f"DOCX error: {exc}")

    # Report text standalone
    txt_path = f"{output_dir}/{ts}_boao_deep_analysis.txt"
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(report_text)
    print(f"TXT: {txt_path}")

    # P1-5: Record trend snapshot
    try:
        trend_path = record_snapshot("boao_forum_2026", articles, dist, clusters)
        print(f"Trend: {trend_path}")
    except Exception as exc:
        print(f"Trend snapshot failed: {exc}")

    print(f"\nDone! Report: {len(report_text)} chars, {len(clusters)} clusters")


if __name__ == "__main__":
    asyncio.run(main())
